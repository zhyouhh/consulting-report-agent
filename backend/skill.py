from datetime import datetime
from pathlib import Path
import json
import logging
import math
import mimetypes
import os
import re
import shutil
import tempfile
import unicodedata
import uuid
from typing import Iterable, Optional

logger = logging.getLogger(__name__)


class StaleFileError(Exception):
    """Raised when a user write's base_mtime_ns no longer matches the file on disk (an AI
    write or another save landed in between). Carries the current mtime_ns (str) for 409."""

    def __init__(self, current_mtime_ns: str):
        super().__init__("文件已被更新")
        self.current_mtime_ns = current_mtime_ns


class UserWriteForbiddenError(Exception):
    """Raised when a user write targets a file outside the editable whitelist — a DOMAIN
    decision, deliberately NOT the built-in PermissionError. os.replace()/write_text() also
    raise PermissionError when the target is locked by an external program (Word/OneDrive/AV
    on Windows); the endpoint must tell that retryable OS failure (→ 500) apart from
    'this file is not user-editable' (→ 403). (codex backend review BLOCKER)"""


class SkillEngine:
    """咨询技能工作流引擎"""

    CORE_CONTEXT_FILES = [
        ("当前项目概览", "plan/project-overview.md"),
        ("当前项目进度", "plan/progress.md"),
        ("阶段门禁", "plan/stage-gates.md"),
        ("项目备注", "plan/notes.md"),
    ]

    FORMAL_PLAN_FILES = {
        "project-overview.md",
        "progress.md",
        "stage-gates.md",
        "notes.md",
        "outline.md",
        "research-plan.md",
        "references.md",
        "tasks.md",
        "review.md",
        "data-log.md",
        "analysis-notes.md",
        "independent-review.md",
        "lint-report.md",
        "presentation-plan.md",
        "delivery-log.md",
    }

    # R3: 文件语义单一真值源。键为「完整相对 posix 路径」（非 basename）——否则
    # materials/imported/outline.md 会被误判为 S1 正式大纲。stage 是文件级属性（用于置顶）。
    FILE_SEMANTICS = {
        "plan/project-overview.md": {"group": "overview", "stage": "S0"},
        "plan/notes.md": {"group": "research", "stage": "S1"},
        "plan/references.md": {"group": "research", "stage": "S1"},
        "plan/data-log.md": {"group": "research", "stage": "S2"},
        "plan/outline.md": {"group": "analysis", "stage": "S1"},
        "plan/research-plan.md": {"group": "analysis", "stage": "S1"},
        "plan/analysis-notes.md": {"group": "analysis", "stage": "S3"},
        "content/report_draft_v1.md": {"group": "draft", "stage": "S4"},
        "plan/independent-review.md": {"group": "review", "stage": "S5"},
        "plan/lint-report.md": {"group": "review", "stage": "S5"},
        "plan/presentation-plan.md": {"group": "delivery", "stage": "S6"},
        "plan/delivery-log.md": {"group": "delivery", "stage": "S7"},
        "plan/stage-gates.md": {"group": "tracking", "stage": None},
        "plan/progress.md": {"group": "tracking", "stage": None},
        "plan/tasks.md": {"group": "tracking", "stage": None},
        "plan/review.md": {"group": "other", "stage": None},
    }

    # R3: 用户可手动编辑白名单（canonical = casefold 后的完整 posix 相对路径）。默认 deny——
    # 任何不在此集合的文件（后端自动维护 / 审查报告 / 退役 / checkpoint）都只读。
    USER_EDITABLE_FILES = {
        "content/report_draft_v1.md",
        "plan/outline.md",
        "plan/research-plan.md",
        "plan/notes.md",
        "plan/references.md",
        "plan/data-log.md",
        "plan/analysis-notes.md",
        "plan/presentation-plan.md",
    }

    # R3: GET /files 跳过的退役文件（不显示）。
    RETIRED_WORKSPACE_FILES = {
        "plan/project-info.md",
        "plan/review-checklist.md",
    }

    STAGE_CHECKPOINTS_FILENAME = "stage_checkpoints.json"
    STAGE_CHECKPOINT_KEYS = {
        "s0_interview_done_at",
        "outline_confirmed_at",
        "review_started_at",
        "review_passed_at",
        "presentation_ready_at",
        "delivery_archived_at",
    }
    MIGRATION_MARKER_KEY = "__migrated_at"
    METHODOLOGY_SNAPSHOT_KEY = "__methodology_snapshot"
    # 非 checkpoint 的受保护内部 string 键集合（确认时快照的方法论 + migration marker）。
    # 绝不加进 STAGE_CHECKPOINT_KEYS——那个有 `set(_CASCADE_ORDER) == STAGE_CHECKPOINT_KEYS`
    # 的 invariant assert（:117），加了即炸（红队 R3）。_load_stage_checkpoints 只返回
    # STAGE_CHECKPOINT_KEYS 的 str，故这些键天然不经 get_workspace_summary 暴露给前端。
    PRESERVED_STAGE_CHECKPOINT_STRING_KEYS = {MIGRATION_MARKER_KEY, METHODOLOGY_SNAPSHOT_KEY}
    assert not (PRESERVED_STAGE_CHECKPOINT_STRING_KEYS & STAGE_CHECKPOINT_KEYS), (
        "preserved string keys must never overlap STAGE_CHECKPOINT_KEYS"
    )
    _CASCADE_ORDER = [
        "s0_interview_done_at",
        "outline_confirmed_at",
        "review_started_at",
        "review_passed_at",
        "presentation_ready_at",
        "delivery_archived_at",
    ]
    assert set(_CASCADE_ORDER) == STAGE_CHECKPOINT_KEYS, (
        "_CASCADE_ORDER must cover exactly STAGE_CHECKPOINT_KEYS"
    )
    _EXPECTED_LENGTH_LINE_PATTERN = re.compile(r"预期篇幅[^\n]*?[:：]\s*([^\n(（]+)")
    _EXPECTED_LENGTH_HEADING_PATTERN = re.compile(
        r"^##\s*预期篇幅\s*\n\s*([^\n(（]+)",
        re.MULTILINE,
    )
    _DL_ENTRY_PATTERN = re.compile(r"^#{3,4}\s*\*{0,2}\s*\[(DL-[^\]]+)\]", re.MULTILINE)
    _DL_REFERENCE_PATTERN = re.compile(r"\[(DL-[^\]]+)\]")
    _DL_REFERENCE_BRACKET_PATTERN = re.compile(r"\[([^\]]*DL-[^\]]*)\]")
    _DL_REFERENCE_GROUP_PATTERN = re.compile(r"DL-(?:\d{4}-)?\d+(?:/\d+)*")
    _MARKDOWN_STRIP_PATTERNS = [
        (re.compile(r"```[\s\S]*?```"), ""),
        (re.compile(r"`[^`]*`"), ""),
        (re.compile(r"!\[[^\]]*\]\([^)]*\)"), ""),
        (re.compile(r"\[([^\]]*)\]\([^)]*\)"), r"\1"),
        (re.compile(r"^\s{0,3}#{1,6}\s+", re.MULTILINE), ""),
        (re.compile(r"^\s*[-*+]\s+", re.MULTILINE), ""),
        (re.compile(r"^\s*\d+\.\s+", re.MULTILINE), ""),
        (re.compile(r"\*\*([^*]+)\*\*"), r"\1"),
        (re.compile(r"\*([^*]+)\*"), r"\1"),
        (re.compile(r"^\s*\|.*\|\s*$", re.MULTILINE), ""),
    ]
    _EVIDENCE_MARKERS = (
        re.compile(r"https?://"),
        re.compile(r"material:[a-zA-Z0-9\-]+"),
        re.compile(r"^(访谈|调研)[:：]", re.MULTILINE),
    )
    _SELF_SIGNATURE_PATTERNS = [
        re.compile(r"审查人\s*[:：]\s*(咨询报告写作助手|AI|助手|Claude|GPT|ChatGPT|gemini|模型)"),
    ]
    _PREMATURE_REVIEW_VERDICT_PATTERNS = [
        re.compile(r"审查结论\s*[:：]"),
        re.compile(r"建议通过"),
        re.compile(r"审查通过"),
    ]
    _ARCHIVE_CLAIM_PATTERNS = [
        re.compile(r"(项目状态|交付状态)[^\n]*?[:：]?\s*(已完成|已交付|已归档|已结束)"),
    ]
    # Broadened from spec's `客户反馈` to `反馈` to catch variants like `**反馈 A**`
    # that the literal spec regex would miss. Covers all feedback headings, not only
    # customer-facing ones, which also closes a real bypass vector.
    _DELIVERY_PLACEHOLDER_INLINE = re.compile(
        r"-\s*\[x\][^\n]*反馈[^\n]*[(（]?\s*(待记录|待补充|暂无)\s*[)）]?"
    )
    _DELIVERY_BLOCK_RE = re.compile(
        r"-\s*\[x\][^\n]*反馈[^\n]*\n(?P<body>(?:[^\n]*(?:\n|$)){0,5})",
        re.MULTILINE,
    )
    _PLACEHOLDER_WORDS_RE = re.compile(r"[(（]?\s*(待记录|待补充|暂无)\s*[)）]?")
    IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp"}
    TEXT_SUFFIXES = {".md", ".txt", ".csv"}
    STAGE_CHECKLIST_ITEMS = {
        "S0": [
            "\u9700\u6c42\u8bbf\u8c08\u5b8c\u6210",
            "\u8303\u56f4\u754c\u5b9a\u660e\u786e",
            "project-overview.md \u521b\u5efa",
            "\u4ea4\u4ed8\u5f62\u5f0f\u786e\u8ba4",
        ],
        "S1": [
            "notes.md \u66f4\u65b0",
            "references.md \u66f4\u65b0",
            "\u5206\u6790\u6846\u67b6\u786e\u5b9a",
            "outline.md \u5b8c\u6210",
            "research-plan.md \u5b8c\u6210",
        ],
        "S2": [
            "data-log.md \u66f4\u65b0",
            "\u4e00\u624b/\u4e8c\u624b\u8d44\u6599\u6761\u76ee\u5f55\u5165",
            "\u8bbf\u8c08\u6216\u8c03\u7814\u8bb0\u5f55\u6c89\u6dc0",
        ],
        "S3": [
            "analysis-notes.md \u521b\u5efa/\u66f4\u65b0",
            "\u5173\u952e\u53d1\u73b0\u63d0\u70bc",
            "\u7ed3\u8bba\u4e0e\u5047\u8bbe\u533a\u5206\u6e05\u695a",
        ],
        "S4": [
            "\u62a5\u544a\u7ed3\u6784\u786e\u5b9a",
            "content/report_draft_v1.md \u5f62\u6210\u6709\u6548\u8349\u7a3f",
            "\u5404\u7ae0\u8282\u5185\u5bb9\u6301\u7eed\u5b8c\u5584",
            "\u6267\u884c\u6458\u8981\u4e0e\u56fe\u8868\u540c\u6b65\u66f4\u65b0",
        ],
        "S5": [
            "独立审查完成",
            "AI 味自查完成",
            "\u4e8b\u5b9e\u3001\u903b\u8f91\u4e0e\u8bed\u8a00\u8d28\u91cf\u5ba1\u67e5\u5b8c\u6210",
        ],
        "S6": [
            "\u4ec5\u5f53 project-overview.md \u4e2d\u4ea4\u4ed8\u5f62\u5f0f = \u62a5\u544a+\u6f14\u793a \u65f6\u542f\u7528",
            "presentation-plan.md \u5b8c\u6210",
            "PPT / \u8bb2\u7a3f / Q&A \u51c6\u5907",
        ],
        "S7": [
            "delivery-log.md \u66f4\u65b0",
            "\u5ba2\u6237\u53cd\u9988\u6536\u96c6",
            "\u540e\u7eed\u52a8\u4f5c\u4e0e\u5f52\u6863\u8bb0\u5f55",
        ],
    }
    STAGE_ORDER = ("S0", "S1", "S2", "S3", "S4", "S5", "S6", "S7")
    STAGE_TITLES = {
        "S0": "项目启动",
        "S1": "研究设计",
        "S2": "资料采集",
        "S3": "分析沉淀",
        "S4": "报告撰写",
        "S5": "质量审查",
        "S6": "演示准备",
        "S7": "交付归档",
    }

    # R5: project_type(slug) → modules 文件名。management-document 的 slug 与文件名不一致
    # （文件是 management-system.md），其余 5 个同名。load_type_skeleton 用它定位骨架模块。
    TYPE_SKELETON_MAP = {
        "strategy-consulting": "strategy-consulting.md",
        "market-research": "market-research.md",
        "specialized-research": "specialized-research.md",
        "management-document": "management-system.md",
        "implementation-plan": "implementation-plan.md",
        "due-diligence": "due-diligence.md",
        "technical-bid": "technical-bid.md",
    }

    # R5: 类型→声明腔调（§7.3）。analytical=招牌框架；structural=结构纪律；specialized=按子题。
    METHODOLOGY_TONE = {
        "strategy-consulting": "analytical",
        "market-research": "analytical",
        "due-diligence": "analytical",
        "management-document": "structural",
        "implementation-plan": "structural",
        "specialized-research": "specialized",
        "technical-bid": "bid",
    }

    # R5: 共享分析框架菜单（横向对所有类型可用，v1 仅菜单一行；细节全文留 v2）。
    # 常驻 S1–S4 注入。token 由 test_build_methodology_block_token_budget 实测 ≤2k/轮。
    FRAMEWORK_MENU = (
        "## 可选分析框架菜单（按报告实际需要挑，不被类型锁死；也可用你自己知道的其他框架）\n"
        "- SWOT：内外部优劣势/机会/威胁（广谱）\n"
        "- PEST：政治/经济/社会/技术宏观环境（广谱·战略）\n"
        "- 波特五力：行业竞争强度五维（战略/市场/尽调）\n"
        "- 价值链：主要+支持活动定位优势环节（战略）\n"
        "- 金字塔原理/MECE：结论先行、不重不漏分组（广谱）\n"
        "- 对标分析：选可比对象横向比（广谱）\n"
        "- 根因分析：问题溯源不停表面（专项研究）\n"
        "- 成熟度模型：五级阶梯定位现状/目标（评估类）\n"
        "- BCG/GE 矩阵：业务组合定位（战略）\n"
        "- 安索夫矩阵：增长路径四象限（战略）\n"
        "- TAM-SAM-SOM：市场规模自上而下（市场）\n"
        "- CR4/HHI：市场集中度（市场）\n"
        "- SMART：目标设定五要素（实施方案）\n"
        "- RACI：责任分配四角色（实施方案）\n"
        "- 甘特/里程碑：进度与关键节点（实施方案）\n"
        "- 财务尽调三维：收入真实性/成本/资产质量（尽调）\n"
        "- 红旗识别：异常/诉讼/关联交易（尽调）\n"
        "- 影响-可行矩阵：建议优先级排序（广谱·建议）\n"
        "- DAMA-DMBOK / ISO 8000：数据治理组织/质量/成熟度（数据专项）\n"
    )

    # R5: 声明行格式（行首关键词 + 顿号/逗号分隔；既给人看又可解析，不用隐藏 marker）。
    _METHODOLOGY_DECLARATION_RE = re.compile(
        r"^[^\S\n]*\*{0,2}方法论框架\*{0,2}[^\S\n]*[:：][^\S\n]*(.+?)[^\S\n]*$",
        re.MULTILINE,
    )
    # 精确匹配放行的已知框架名（**无空格 casefold**，比对时把 token 也去空格归一化，
    # 让「BCG 矩阵」「ISO 8000」带空格写法也命中）。与 FRAMEWORK_MENU 并行维护：菜单是给
    # 模型看的一句话清单，这里是给净化用的精确名集。
    KNOWN_FRAMEWORK_NAMES = {
        "swot", "pest", "波特五力", "五力", "价值链", "金字塔原理", "金字塔原理/mece",
        "mece", "金字塔", "对标分析", "根因分析", "成熟度模型", "bcg", "bcg矩阵",
        "bcg/ge矩阵", "ge矩阵", "安索夫矩阵", "tam-sam-som", "cr4", "hhi", "cr4/hhi",
        "smart", "raci", "甘特", "里程碑", "甘特/里程碑", "财务尽调三维", "红旗识别",
        "影响-可行矩阵", "dama-dmbok", "iso8000", "dama-dmbok/iso8000", "章-条-款-项",
    }
    # 原样 casefold 子串命中即整条 malformed：注入符号 / 中文操控·控制语义词。
    # 工具名 + checkpoint + 英文操作词移到 _METHODOLOGY_DANGER_NORMALIZED（归一化匹配，防分隔符/拆词绕过）。
    _METHODOLOGY_DANGER_SUBSTRINGS = (
        "__", "<stage", "stage-ack", "ignore",
        "系统提示", "系統提示", "忽略", "覆写", "覆寫", "覆盖",
        # 中文阶段操控 / 注入指令 / 控制语义词（不含框架·业务常用词：分析/模型/矩阵/交付/阶段）
        "推进", "回退", "归档", "无视", "跳过", "停止", "立即", "删除",
        "设为", "标记为", "指令", "请你", "门禁", "检查点",
    )
    # 归一化（NFKC+casefold+删 Cf 格式字符+去空白/分隔符）后子串命中即 malformed：工具名 +
    # 全部 6 个 STAGE_CHECKPOINT_KEYS 的去分隔符形态 + 英文操作词。防 "advance stage" /
    # "advance-stage" / "s0 interview done at" / "over ride" 等变体绕过（红队 v2/v3、quality NIT）。
    # 零误杀（真框架名不会等于这些 API 标识符）。守护测试 test_*_all_checkpoint_key_variants 锁
    # checkpoint 全覆盖，防未来加 key 漏项。
    _METHODOLOGY_DANGER_NORMALIZED = (
        # 工具名
        "writefile", "editfile", "appendreportdraft", "readfile", "readmaterialfile",
        "websearch", "fetchurl",
        # 阶段机 + 全部 6 个 checkpoint 的去分隔符形态（含 v3 漏掉的 s0interviewdone）
        "advancestage", "checkpoint", "stageack", "s0interviewdone",
        "outlineconfirmed", "reviewstarted", "reviewpassed", "presentationready",
        "deliveryarchived",
        # 英文操作词（归一化防 "over ride" 拆词，quality NIT）
        "override", "prompt",
    )

    REPORT_DRAFT_PATH = "content/report_draft_v1.md"
    REPORT_DRAFT_CANDIDATES = (REPORT_DRAFT_PATH,)
    INDEPENDENT_REVIEW_ANCHORS = (
        "## 1. 结论-证据一致性",
        "## 2. 关键假设与逻辑链",
        "## 3. 数据口径一致性",
        "## 4. 建议可执行性",
        "## 5. 目标读者匹配",
    )
    INDEPENDENT_REVIEW_COMPLETION_MARKER = "<!-- independent-review:complete -->"
    LINT_REPORT_ANCHORS = ("## 按章节排列", "## 总览")
    LINT_REPORT_COMPLETION_MARKER = "<!-- lint-report:complete -->"
    CHECKPOINT_PREREQ = {
        "s0_interview_done_at": None,
        "outline_confirmed_at": (
            "_has_effective_outline",
            "plan/outline.md",
            "需要先生成有效报告大纲，才能确认大纲并进入资料采集。",
            "请先让助手补齐 `plan/outline.md`，再确认大纲。",
        ),
        "review_started_at": (
            "_has_effective_report_draft",
            REPORT_DRAFT_PATH,
            "需要先形成有效报告正文，才能进入质量审查。",
            f"请先让助手写入 `{REPORT_DRAFT_PATH}`，再开始审查。",
        ),
        "review_passed_at": (
            "_has_effective_review_reports",
            "plan/independent-review.md, plan/lint-report.md",
            "需要先完成独立审查和 AI 味自查，才能标记审查通过。",
            "请先在 S5 阶段点击上方'独立审查'和'AI 味自查'按钮，再确认审查通过。",
        ),
        "presentation_ready_at": (
            "_has_effective_presentation_plan",
            "plan/presentation-plan.md",
            "需要先完成有效演示方案，才能标记演示准备完成。",
            "请先让助手补齐 `plan/presentation-plan.md`，再确认演示准备完成。",
        ),
        "delivery_archived_at": (
            "_has_effective_delivery_log",
            "plan/delivery-log.md",
            "需要先记录有效交付归档信息，才能结束项目。",
            "请先让助手补齐 `plan/delivery-log.md`，再归档结束项目。",
        ),
    }

    def _stage_checkpoints_path(self, project_path):
        return Path(project_path) / self.STAGE_CHECKPOINTS_FILENAME

    def _load_stage_checkpoints(self, project_path) -> dict[str, str]:
        raw = self._read_raw_stage_checkpoints(project_path)
        return {
            key: value
            for key, value in raw.items()
            if key in self.STAGE_CHECKPOINT_KEYS and isinstance(value, str)
        }

    def _read_raw_stage_checkpoints(self, project_path) -> dict:
        checkpoints_path = self._stage_checkpoints_path(project_path)
        if not checkpoints_path.exists():
            return {}
        try:
            data = json.loads(checkpoints_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError, TypeError, ValueError):
            return {}
        if not isinstance(data, dict):
            return {}
        return data

    def _write_raw_stage_checkpoints(self, project_path, data):
        checkpoints_path = self._stage_checkpoints_path(project_path)
        checkpoints_path.parent.mkdir(parents=True, exist_ok=True)
        payload = json.dumps(data, ensure_ascii=False, indent=2)
        # 原子写：同目录唯一 temp（mkstemp，避免固定 temp 名被并发 writer 抢占，对齐
        # user_write_file 模式）+ os.replace；异常清理自己的 temp（codex B4 quality/红队）。
        tmp_fd, tmp_name = tempfile.mkstemp(dir=str(checkpoints_path.parent), suffix=".tmp")
        try:
            with os.fdopen(tmp_fd, "w", encoding="utf-8") as handle:
                handle.write(payload)
            os.replace(tmp_name, checkpoints_path)
        except BaseException:
            try:
                os.unlink(tmp_name)
            except OSError:
                pass
            raise

    def _backfill_stage_checkpoints_if_missing(self, project_path):
        """Schema-incremental migration for stage_checkpoints.json.

        Runs idempotently on every project load. Scenarios:
        1. File missing → create with __migrated_at; backfill by stage-gates.md
        2. File exists, missing s0_interview_done_at:
           - stage-gates.md shows stage ≥ S1, OR any downstream checkpoint set
             → backfill s0
           - otherwise (stage=S0, no downstream) → do NOT backfill
        3. Key already present → no-op
        4. outline_confirmed_at still backfills only at stage ≥ S2 (unchanged)
        """
        checkpoints_path = self._stage_checkpoints_path(project_path)
        stage_gates_path = Path(project_path) / "plan" / "stage-gates.md"
        timestamp = datetime.now().isoformat(timespec="seconds")

        raw = self._read_raw_stage_checkpoints(project_path)
        file_existed_before = checkpoints_path.exists()
        changed = False

        if not file_existed_before:
            raw = {self.MIGRATION_MARKER_KEY: timestamp}
            changed = True

        current_stage = None
        if stage_gates_path.exists():
            stage_text = stage_gates_path.read_text(encoding="utf-8")
            current_stage = self._extract_stage_code(stage_text)

        # Downstream = any checkpoint other than s0
        has_downstream = any(
            key in raw
            for key in self._CASCADE_ORDER
            if key != "s0_interview_done_at"
        )

        # Backfill s0 when stage >= S1 OR downstream present
        if "s0_interview_done_at" not in raw:
            stage_ok = False
            if current_stage:
                try:
                    stage_ok = self._stage_index(current_stage) >= self._stage_index("S1")
                except ValueError:
                    stage_ok = False  # malformed stage-gates; stay cautious
            if stage_ok or has_downstream:
                raw["s0_interview_done_at"] = timestamp
                changed = True

        # Preserve legacy outline backfill (stage >= S2)
        if "outline_confirmed_at" not in raw and current_stage:
            try:
                if self._stage_index(current_stage) >= self._stage_index("S2"):
                    raw["outline_confirmed_at"] = timestamp
                    changed = True
            except ValueError:
                pass

        if changed:
            # 写回前重读 PRESERVED 键合并，避免覆盖并发写入的 snapshot/marker（红队 B4 BLOCKER 3
            # 纵深防御；注：确认大纲那刻 backfill 通常 changed=False[s0 已在、stage<S2 不补 outline]，
            # 实际并发窗口窄）。backfill 不改 PRESERVED 键，重读最新值合并即可。
            latest = self._read_raw_stage_checkpoints(project_path)
            for preserved_key in self.PRESERVED_STAGE_CHECKPOINT_STRING_KEYS:
                if preserved_key in latest:
                    raw[preserved_key] = latest[preserved_key]
            self._write_raw_stage_checkpoints(project_path, raw)

    def _clear_stage_checkpoint_cascade(self, project_path, key):
        if key not in self._CASCADE_ORDER:
            raise ValueError(f"unknown cascade key: {key}")

        start = self._CASCADE_ORDER.index(key)
        checkpoints = self._load_stage_checkpoints(project_path)
        changed = False
        for cascade_key in self._CASCADE_ORDER[start:]:
            if cascade_key in checkpoints:
                del checkpoints[cascade_key]
                changed = True

        if changed:
            raw = self._read_raw_stage_checkpoints(project_path)
            marker = raw.get(self.MIGRATION_MARKER_KEY)
            payload = dict(checkpoints)
            if marker:
                payload[self.MIGRATION_MARKER_KEY] = marker
            # R5: 方法论快照仅当被清范围**不含** outline_confirmed_at 时保留（红队 R2）。
            # 清 outline_confirmed_at 本身（或上游 s0）→ 删快照；清 review_*/下游 → 保留，
            # 否则用户 S5「回去改」会丢快照、S2–S4 退回 default。
            snapshot = raw.get(self.METHODOLOGY_SNAPSHOT_KEY)
            if (
                isinstance(snapshot, str)
                and "outline_confirmed_at" not in self._CASCADE_ORDER[start:]
            ):
                payload[self.METHODOLOGY_SNAPSHOT_KEY] = snapshot
            self._write_raw_stage_checkpoints(project_path, payload)

    def _save_stage_checkpoint(self, project_path, key):
        if key not in self.STAGE_CHECKPOINT_KEYS:
            raise ValueError(f"Unsupported stage checkpoint key: {key}")

        raw = self._read_raw_stage_checkpoints(project_path)
        existing = raw.get(key)
        if isinstance(existing, str):
            return existing

        timestamp = datetime.now().isoformat(timespec="seconds")
        raw[key] = timestamp
        self._write_raw_stage_checkpoints(project_path, raw)
        return timestamp

    def _clear_stage_checkpoint(self, project_path, key):
        if key not in self.STAGE_CHECKPOINT_KEYS:
            raise ValueError(f"Unsupported stage checkpoint key: {key}")

        raw = self._read_raw_stage_checkpoints(project_path)
        if key not in raw:
            return

        raw.pop(key, None)
        self._write_raw_stage_checkpoints(project_path, raw)

    def get_stage_checkpoint_prereq_notice(self, key: str) -> dict | None:
        prereq = self.CHECKPOINT_PREREQ.get(key)
        if not prereq:
            return None
        _, path, reason, user_action = prereq
        return {"path": path, "reason": reason, "user_action": user_action}

    def _validate_stage_checkpoint_prereq(self, project_path: Path, key: str) -> None:
        prereq = self.CHECKPOINT_PREREQ.get(key)
        if not prereq:
            return
        validator_name, path, reason, _ = prereq
        validator = getattr(self, validator_name)
        if not validator(project_path):
            raise ValueError(f"{reason} 缺少前置文件: {path}")
        self._validate_stage_checkpoint_predecessors(project_path, key)

    def _required_stage_checkpoint_predecessors(self, project_path: Path, key: str) -> tuple[str, ...]:
        if key not in self._CASCADE_ORDER:
            return ()
        required = list(self._CASCADE_ORDER[:self._CASCADE_ORDER.index(key)])
        if key == "delivery_archived_at" and not self._delivery_mode_requires_presentation(project_path):
            required = [item for item in required if item != "presentation_ready_at"]
        return tuple(required)

    def _validate_stage_checkpoint_predecessors(self, project_path: Path, key: str) -> None:
        required = self._required_stage_checkpoint_predecessors(project_path, key)
        if not required:
            return
        checkpoints = self._load_stage_checkpoints(project_path)
        missing = [item for item in required if item not in checkpoints]
        if missing:
            raise ValueError(f"缺少前序阶段 checkpoint: {', '.join(missing)}")

    def _stage_one_completion_state(self, project_path: Path, checkpoints: dict[str, str] | None = None) -> dict:
        checkpoints = checkpoints if checkpoints is not None else self._load_stage_checkpoints(project_path)

        project_overview_ready = self._is_effective_plan_file(project_path, "project-overview.md")
        notes_ready = self._has_effective_notes(project_path)
        references_ready = self._has_effective_references(project_path)
        outline_ready = self._has_effective_outline(project_path)
        research_plan_ready = self._has_effective_research_plan(project_path)
        interview_done = "s0_interview_done_at" in checkpoints
        outline_confirmed = "outline_confirmed_at" in checkpoints

        missing_prerequisites = []
        if not project_overview_ready:
            missing_prerequisites.append("project-overview.md")
        if not interview_done:
            missing_prerequisites.append("S0 需求访谈")
        if not notes_ready:
            missing_prerequisites.append("notes.md")
        if not references_ready:
            missing_prerequisites.append("references.md")
        if not outline_ready:
            missing_prerequisites.append("outline.md")
        if not research_plan_ready:
            missing_prerequisites.append("research-plan.md")

        missing_for_completion = list(missing_prerequisites)
        if not outline_confirmed:
            missing_for_completion.append("outline_confirmed_at")

        return {
            "project_overview_ready": project_overview_ready,
            "notes_ready": notes_ready,
            "references_ready": references_ready,
            "outline_ready": outline_ready,
            "research_plan_ready": research_plan_ready,
            "interview_done": interview_done,
            "outline_confirmed": outline_confirmed,
            "stage_zero_complete": project_overview_ready and interview_done,
            "stage_one_prerequisites_complete": not missing_prerequisites,
            "stage_one_complete": not missing_for_completion,
            "missing_prerequisites": missing_prerequisites,
            "missing_for_completion": missing_for_completion,
        }

    def _stage_four_completion_state(
        self,
        project_path: Path,
        checkpoints: dict[str, str] | None = None,
        targets: dict | None = None,
        stage_one_state: dict | None = None,
    ) -> dict:
        checkpoints = checkpoints if checkpoints is not None else self._load_stage_checkpoints(project_path)
        targets = targets if targets is not None else self._resolve_length_targets(project_path)
        stage_one_state = stage_one_state or self._stage_one_completion_state(project_path, checkpoints)

        data_ready = self._has_enough_data_log_sources(project_path, targets["data_log_min"])
        analysis_ready = self._has_enough_analysis_refs(project_path, targets["analysis_refs_min"])
        report_ready = self._has_effective_report_draft(project_path, min_words=targets["report_word_floor"])
        review_started = "review_started_at" in checkpoints

        missing_for_review_start = list(stage_one_state["missing_for_completion"])
        if not data_ready:
            missing_for_review_start.append("data-log.md")
        if not analysis_ready:
            missing_for_review_start.append("analysis-notes.md")
        if not report_ready:
            missing_for_review_start.append(self.REPORT_DRAFT_PATH)

        missing_for_stage_four = list(missing_for_review_start)
        if not review_started:
            missing_for_stage_four.append("review_started_at")

        stage_two_complete = stage_one_state["stage_one_complete"] and data_ready
        stage_three_complete = stage_two_complete and analysis_ready

        return {
            "data_log_quality_ok": data_ready,
            "analysis_quality_ok": analysis_ready,
            "report_ready": report_ready,
            "review_started": review_started,
            "stage_two_complete": stage_two_complete,
            "stage_three_complete": stage_three_complete,
            "review_start_prerequisites_complete": not missing_for_review_start,
            "stage_four_complete": not missing_for_stage_four,
            "missing_for_review_start": missing_for_review_start,
            "missing_for_stage_four": missing_for_stage_four,
        }

    def _stage_five_completion_state(
        self,
        project_path: Path,
        checkpoints: dict[str, str] | None = None,
        targets: dict | None = None,
        stage_one_state: dict | None = None,
        stage_four_state: dict | None = None,
    ) -> dict:
        checkpoints = checkpoints if checkpoints is not None else self._load_stage_checkpoints(project_path)
        targets = targets if targets is not None else self._resolve_length_targets(project_path)
        stage_one_state = stage_one_state or self._stage_one_completion_state(project_path, checkpoints)
        stage_four_state = stage_four_state or self._stage_four_completion_state(
            project_path, checkpoints, targets, stage_one_state
        )

        independent_review_ready = self._has_effective_independent_review(project_path)
        lint_report_ready = self._has_effective_lint_report(project_path)
        review_reports_ready = independent_review_ready and lint_report_ready
        review_passed = "review_passed_at" in checkpoints

        missing_for_review_pass = list(stage_four_state["missing_for_stage_four"])
        if not independent_review_ready:
            missing_for_review_pass.append("independent-review.md（请先点'独立审查'按钮）")
        if not lint_report_ready:
            missing_for_review_pass.append("lint-report.md（请先点'AI 味自查'按钮）")

        missing_for_stage_five = list(missing_for_review_pass)
        if not review_passed:
            missing_for_stage_five.append("review_passed_at")

        return {
            "review_checklist_ready": False,
            "independent_review_ready": independent_review_ready,
            "lint_report_ready": lint_report_ready,
            "review_reports_ready": review_reports_ready,
            "review_passed": review_passed,
            "review_pass_prerequisites_complete": not missing_for_review_pass,
            "stage_five_complete": not missing_for_stage_five,
            "missing_for_review_pass": missing_for_review_pass,
            "missing_for_stage_five": missing_for_stage_five,
        }

    def _stage_six_completion_state(
        self,
        project_path: Path,
        checkpoints: dict[str, str] | None = None,
        targets: dict | None = None,
        stage_one_state: dict | None = None,
        stage_four_state: dict | None = None,
        stage_five_state: dict | None = None,
    ) -> dict:
        checkpoints = checkpoints if checkpoints is not None else self._load_stage_checkpoints(project_path)
        targets = targets if targets is not None else self._resolve_length_targets(project_path)
        stage_one_state = stage_one_state or self._stage_one_completion_state(project_path, checkpoints)
        stage_four_state = stage_four_state or self._stage_four_completion_state(
            project_path, checkpoints, targets, stage_one_state
        )
        stage_five_state = stage_five_state or self._stage_five_completion_state(
            project_path, checkpoints, targets, stage_one_state, stage_four_state
        )

        presentation_required = self._delivery_mode_requires_presentation(project_path)
        presentation_ready = self._has_effective_presentation_plan(project_path)
        presentation_done = "presentation_ready_at" in checkpoints

        missing_for_presentation_ready = list(stage_five_state["missing_for_stage_five"])
        if not presentation_ready:
            missing_for_presentation_ready.append("presentation-plan.md")

        missing_for_stage_six = list(missing_for_presentation_ready)
        if presentation_required and not presentation_done:
            missing_for_stage_six.append("presentation_ready_at")

        return {
            "presentation_required": presentation_required,
            "presentation_ready": presentation_ready,
            "presentation_done": presentation_done,
            "presentation_ready_prerequisites_complete": presentation_required and not missing_for_presentation_ready,
            "stage_six_complete": (
                stage_five_state["stage_five_complete"]
                and ((presentation_ready and presentation_done) if presentation_required else True)
            ),
            "missing_for_presentation_ready": missing_for_presentation_ready,
            "missing_for_stage_six": missing_for_stage_six,
        }

    def _validate_stage_checkpoint_transition(self, project_path: Path, key: str) -> None:
        if key not in self.STAGE_CHECKPOINT_KEYS:
            raise ValueError(f"Unsupported stage checkpoint key: {key}")

        checkpoints = self._load_stage_checkpoints(project_path)
        targets = self._resolve_length_targets(project_path)
        stage_one_state = self._stage_one_completion_state(project_path, checkpoints)
        stage_four_state = self._stage_four_completion_state(project_path, checkpoints, targets, stage_one_state)
        stage_five_state = self._stage_five_completion_state(
            project_path, checkpoints, targets, stage_one_state, stage_four_state
        )
        stage_six_state = self._stage_six_completion_state(
            project_path, checkpoints, targets, stage_one_state, stage_four_state, stage_five_state
        )

        def require(condition: bool, message: str) -> None:
            if not condition:
                raise ValueError(message)

        project_overview_ready = stage_one_state["project_overview_ready"]

        if key == "s0_interview_done_at":
            require(project_overview_ready, "需要先创建有效 project-overview.md，才能完成需求访谈。")
            return

        if key == "outline_confirmed_at":
            missing = stage_one_state["missing_prerequisites"]
            require(not missing, f"需要先补齐 {', '.join(missing)}，才能确认大纲。")
            # R5: 方法论声明前置——仅首次确认（outline_confirmed_at 未 set）+ 已知 6-slug 时校验。
            # 不进 _stage_one_completion_state 持久完成态（否则 legacy 已确认无声明项目被拉回 S1，
            # 红队 BLOCKER 1）；未知 type 不卡（避死锁）。
            if "outline_confirmed_at" not in checkpoints:
                project_type = self._get_project_type_for_path(project_path)
                if project_type in self.TYPE_SKELETON_MAP:
                    outline_text = self._read_plan_file(project_path, "outline.md") or ""
                    state, _ = self.parse_and_sanitize_methodology(outline_text)
                    require(
                        state == "parsed",
                        "大纲缺少有效方法论声明行（如「方法论框架：SWOT、波特五力」），"
                        "请在大纲顶部补一行后再确认。",
                    )
            return

        if key == "review_started_at":
            missing = stage_four_state["missing_for_review_start"]
            require(not missing, f"需要先补齐 {', '.join(missing)}，才能开始审查。")
            return

        if key == "review_passed_at":
            missing = stage_five_state["missing_for_review_pass"]
            require(not missing, f"需要先补齐 {', '.join(missing)}，才能标记审查通过。")
            return

        if key == "presentation_ready_at":
            require(stage_six_state["presentation_required"], "仅报告项目不需要演示准备阶段。")
            missing = stage_six_state["missing_for_presentation_ready"]
            require(not missing, f"需要先补齐 {', '.join(missing)}，才能标记演示准备完成。")
            return

        if key == "delivery_archived_at":
            if stage_six_state["presentation_required"]:
                missing = stage_six_state["missing_for_stage_six"]
            else:
                missing = stage_five_state["missing_for_stage_five"]
            require(not missing, f"需要先补齐 {', '.join(missing)}，才能归档。")
            require(self._has_effective_delivery_log(project_path), "需要先完成 delivery-log.md，才能归档。")

    def _resolve_length_targets(self, project_path):
        overview_path = project_path / "plan" / "project-overview.md"
        expected = 3000
        fallback_used = True
        if overview_path.exists():
            text = overview_path.read_text(encoding="utf-8")
            for pattern in (
                self._EXPECTED_LENGTH_LINE_PATTERN,
                self._EXPECTED_LENGTH_HEADING_PATTERN,
            ):
                match = pattern.search(text)
                if not match:
                    continue
                nums = re.findall(r"\d+", match.group(1))
                if nums:
                    expected = max(int(n) for n in nums)
                    fallback_used = False
                    break
        # Plan §9.3: minimum data-log entries scale with expected report length.
        data_log_min = min(12, math.ceil(expected / 1000 * 1.3))
        # Plan §9.3: analysis citations also scale, capped to keep S3 practical.
        analysis_refs_min = min(8, math.ceil(expected / 1000 * 0.8))
        return {
            "expected_length": expected,
            "data_log_min": max(3, data_log_min),
            "analysis_refs_min": max(2, analysis_refs_min),
            # Plan §9.3: drafts below 70% of target length are not review-ready.
            "report_word_floor": int(expected * 0.7),
            "fallback_used": fallback_used,
        }

    def _count_valid_data_log_sources(self, project_path):
        data_log = project_path / "plan" / "data-log.md"
        if not data_log.exists():
            return 0
        text = data_log.read_text(encoding="utf-8")
        entries = list(self._DL_ENTRY_PATTERN.finditer(text))
        valid = 0
        for idx, match in enumerate(entries):
            start = match.end()
            end = entries[idx + 1].start() if idx + 1 < len(entries) else len(text)
            body = text[start:end]
            if any(pattern.search(body) for pattern in self._EVIDENCE_MARKERS):
                valid += 1
        return valid

    def _has_enough_data_log_sources(self, project_path, min_count):
        return self._count_valid_data_log_sources(project_path) >= min_count

    def _count_analysis_refs(self, project_path):
        analysis = project_path / "plan" / "analysis-notes.md"
        data_log = project_path / "plan" / "data-log.md"
        if not analysis.exists() or not data_log.exists():
            return 0
        refs = self._extract_analysis_ref_ids(
            analysis.read_text(encoding="utf-8"),
            self._extract_data_log_ids(data_log.read_text(encoding="utf-8")),
        )
        return len(refs)

    def _count_analysis_refs_in_text(self, project_path, analysis_text: str) -> int:
        data_log = project_path / "plan" / "data-log.md"
        if not data_log.exists():
            return 0
        refs = self._extract_analysis_ref_ids(
            analysis_text,
            self._extract_data_log_ids(data_log.read_text(encoding="utf-8")),
        )
        return len(refs)

    def _extract_data_log_ids(self, data_log_text: str) -> set[str]:
        return {
            m.group(1)
            for m in self._DL_ENTRY_PATTERN.finditer(data_log_text or "")
        }

    def _extract_analysis_ref_ids(self, analysis_text: str, dl_ids: set[str]) -> set[str]:
        refs: set[str] = set()
        for bracket_match in self._DL_REFERENCE_BRACKET_PATTERN.finditer(analysis_text or ""):
            bracket_content = bracket_match.group(1)
            for group_match in self._DL_REFERENCE_GROUP_PATTERN.finditer(bracket_content):
                refs.update(
                    ref_id
                    for ref_id in self._expand_grouped_dl_ref(group_match.group(0))
                    if ref_id in dl_ids
                )
        return refs

    def _expand_grouped_dl_ref(self, raw_ref: str) -> set[str]:
        parts = (raw_ref or "").split("/")
        first = parts[0].strip()
        if not first:
            return set()
        refs = {first}
        match = re.match(r"^(.*?)(\d+)$", first)
        if not match:
            return refs
        prefix, first_number = match.groups()
        width = len(first_number)
        for suffix in parts[1:]:
            suffix = suffix.strip()
            if not suffix.isdigit():
                continue
            refs.add(f"{prefix}{suffix.zfill(width)}")
        return refs

    def _has_enough_analysis_refs(self, project_path, min_refs):
        return self._count_analysis_refs(project_path) >= min_refs

    def __init__(self, projects_dir: Path, skill_dir: Path):
        self.projects_dir = projects_dir
        self.skill_dir = skill_dir
        self.projects_dir.mkdir(parents=True, exist_ok=True)
        self.registry_path = self.projects_dir / "registry.json"

    def create_project(
        self,
        project_info_or_name=None,
        project_type: str | None = None,
        theme: str | None = None,
        target_audience: str | None = None,
        deadline: str | None = None,
        expected_length: str | None = None,
        notes: str = "",
        workspace_dir: str | None = None,
        initial_material_paths: list[str] | None = None,
        **extra_kwargs,
    ) -> dict:
        """Create new project."""
        payload = self._normalize_create_payload(
            project_info_or_name,
            project_type=project_type,
            theme=theme,
            target_audience=target_audience,
            deadline=deadline,
            expected_length=expected_length,
            notes=notes,
            workspace_dir=workspace_dir,
            initial_material_paths=initial_material_paths,
            **extra_kwargs,
        )

        payload = dict(payload)
        payload["workspace_dir"] = payload.get("workspace_dir") or str(
            self._default_workspace_dir(payload["name"])
        )
        return self._create_workspace_project(payload)

    def _default_workspace_dir(self, project_name: str) -> Path:
        return (self.projects_dir / project_name).resolve()

    def _create_workspace_project(self, payload: dict) -> dict:
        registry = self._load_registry()
        workspace_path = Path(payload["workspace_dir"]).expanduser().resolve()
        if not workspace_path.exists():
            workspace_path.mkdir(parents=True)
        if not workspace_path.is_dir():
            raise ValueError("工作目录无效")
            raise ValueError("宸ヤ綔鐩綍鏃犳晥")

        project_dir = workspace_path / ".consulting-report"
        if project_dir.exists():
            raise ValueError("该工作目录已经初始化过项目")
            raise ValueError("工作目录无效")

        for project in registry["projects"]:
            if Path(project["workspace_dir"]).resolve() == workspace_path:
                raise ValueError("该工作目录已经初始化过项目")

        project_dir.mkdir(parents=True)
        self._initialize_project_structure(project_dir)
        self._populate_v2_plan_files(
            project_path=project_dir,
            name=payload["name"],
            project_type=payload["project_type"],
            theme=payload["theme"],
            target_audience=payload.get("target_audience", ""),
            deadline=payload["deadline"],
            expected_length=payload["expected_length"],
            notes=payload["notes"],
        )

        now = datetime.now().isoformat(timespec="seconds")
        project_record = {
            "id": self._new_id("proj"),
            "name": payload["name"],
            "project_type": payload["project_type"],
            "theme": payload["theme"],
            "target_audience": payload.get("target_audience", ""),
            "deadline": payload["deadline"],
            "expected_length": payload["expected_length"],
            "notes": payload["notes"],
            "workspace_dir": str(workspace_path),
            "project_dir": str(project_dir),
            "created_at": now,
            "updated_at": now,
        }

        registry["projects"].append(project_record)
        self._save_registry(registry)
        self._save_materials(project_record, [])

        if payload["initial_material_paths"]:
            self.add_materials(
                project_record["id"],
                payload["initial_material_paths"],
                added_via="project_create",
            )

        return project_record

    def _initialize_project_structure(self, project_path: Path):
        plan_path = project_path / "plan"
        content_path = project_path / "content"
        output_path = project_path / "output"
        imported_materials_path = project_path / "materials" / "imported"

        plan_path.mkdir(parents=True, exist_ok=True)
        content_path.mkdir(exist_ok=True)
        output_path.mkdir(exist_ok=True)
        imported_materials_path.mkdir(parents=True, exist_ok=True)

        template_dir = self.skill_dir / "plan-template"
        for template_name in sorted(self.FORMAL_PLAN_FILES):
            template_file = template_dir / template_name
            if template_file.exists():
                shutil.copy(template_file, plan_path / template_name)

    def _populate_v2_plan_files(
        self,
        project_path: Path,
        name: str,
        project_type: str,
        theme: str,
        target_audience: str = "",
        deadline: str = "",
        expected_length: str = "",
        notes: str = "",
    ):
        today = datetime.now().strftime("%Y-%m-%d")
        target_audience = (target_audience or "").strip()  # 防御 None：replace 第二参须为 str
        # 目标读者已非必填：留空时项目目标句省略「面向…」前缀，避免病句。
        project_goal = (
            f"面向{target_audience}形成{expected_length}规模的咨询报告初稿"
            if target_audience
            else f"形成{expected_length}规模的咨询报告初稿"
        )
        replacements = {
            "[填写项目名称]": name,
            "[战略咨询/市场研究/专项研究/管理制度/实施方案/尽职调查/技术标]": project_type,
            "[描述客户背景、行业环境、当前面临的挑战]": theme,
            "[具体、可衡量的项目目标]": project_goal,
            "[填写目标读者]": target_audience,
            "[例如：3000字 / 5000-8000字]": expected_length,
            "[填写负责人]": "",
            "[填写报告主题]": theme,
            "[填写客户名称]": "",
        }

        overview_path = project_path / "plan" / "project-overview.md"
        if overview_path.exists():
            content = overview_path.read_text(encoding="utf-8")
            for source, target in replacements.items():
                content = content.replace(source, target)
            content = content.replace("**交付时间**: [YYYY-MM-DD]", f"**交付时间**: {deadline}")
            content = content.replace("**开始日期**: [YYYY-MM-DD]", f"**开始日期**: {today}")
            content = content.replace("**截止日期**: [YYYY-MM-DD]", f"**截止日期**: {deadline}")
            if target_audience and "目标读者" not in content:
                content += f"\n\n## 目标读者\n{target_audience}\n"
            if expected_length and "预期篇幅" not in content:
                content += f"\n## 预期篇幅\n{expected_length}\n"
            overview_path.write_text(content, encoding="utf-8")

        progress_path = project_path / "plan" / "progress.md"
        if progress_path.exists():
            content = progress_path.read_text(encoding="utf-8")
            content = content.replace("[S0/S1/S2/S3/S4/S5/S6/S7]", "S0")
            content = content.replace("[进行中/已完成/待开始/阻塞]", "进行中")
            content = content.replace("[YYYY-MM-DD]", datetime.now().strftime("%Y-%m-%d"), 1)
            progress_path.write_text(content, encoding="utf-8")

        stage_gates_path = project_path / "plan" / "stage-gates.md"
        if stage_gates_path.exists():
            content = stage_gates_path.read_text(encoding="utf-8")
            content = content.replace("**阶段**: [S0/S1/S2/S3/S4/S5/S6/S7]", "**阶段**: S0")
            content = content.replace("**状态**: [进行中/已完成/待开始/阻塞]", "**状态**: 进行中")
            content = content.replace("**更新日期**: [YYYY-MM-DD]", f"**更新日期**: {today}")
            stage_gates_path.write_text(content, encoding="utf-8")

        notes_path = project_path / "plan" / "notes.md"
        if notes_path.exists() and notes:
            content = notes_path.read_text(encoding="utf-8")
            content += f"\n\n## 初始化备注\n{notes}\n"
            notes_path.write_text(content, encoding="utf-8")

    def get_project_record(self, project_ref: str) -> Optional[dict]:
        registry = self._load_registry()
        for project in registry["projects"]:
            if project["id"] == project_ref or project["name"] == project_ref:
                return dict(project)
        return None

    def get_project_path(self, project_ref: str) -> Optional[Path]:
        project_record = self.get_project_record(project_ref)
        if project_record:
            project_path = Path(project_record["project_dir"])
            return project_path if project_path.exists() else None
        return None

    def list_projects(self) -> list:
        registry = self._load_registry()
        return [
            {
                **project,
                "path": project["project_dir"],
            }
            for project in registry["projects"]
        ]

    def list_materials(self, project_ref: str) -> list[dict]:
        project_record = self.get_project_record(project_ref)
        if not project_record:
            return []
        materials = self._load_materials(project_record)
        for material in materials:
            status, reason = self._material_conversion_status(project_record, material)
            material["conversion_status"] = status
            material["conversion_reason"] = reason
        return materials

    def _material_conversion_status(self, project_record: dict, material: dict) -> tuple[str, str | None]:
        """N6 D2：只读探测材料转换状态，供材料列表展示。
        必须健壮（任何缺失/异常一律降级 not_parsed），绝不抛——材料列表接口不能 500。

        N6 Fix3（perf）：状态探测是 advisory 展示，绝不 re-hash live 文件（list 调一次就 O(n×文件大小)）。
        改为用 add-time 存的 content_sha256 算 key（inline 派生，不经 _cache_key_for_material——后者
        会读 live 文件）。常见的 imported/chat_upload 不可变场景下 add-time hash == live hash、状态正确；
        被改过的 workspace 文件 chip 可能略 stale，下次真正 read 时自愈，advisory 可接受。
        """
        converter = getattr(self, "_material_converter", None)
        if converter is None:
            return "not_parsed", None
        content_sha256 = material.get("content_sha256")
        if not content_sha256:
            return "not_parsed", None
        try:
            # 源文件已被删/移走时，旧 .md/.error 缓存仍可能在 → 必须降级 not_parsed，
            # 不能据陈旧缓存误报 parsed/failed。存在性检查是廉价 stat、不 re-hash（守住 Fix3 perf）。
            material_path = self._resolve_material_path(project_record, material)
            if not material_path.exists():
                return "not_parsed", None
            extra = converter.image_cache_extra if material.get("media_kind") == "image_like" else ""
            key = converter.cache_key_from_sha256(content_sha256, extra)
            return converter.status_for_key(key)
        except Exception:  # noqa: BLE001 探测失败一律降级，绝不阻断材料列表
            return "not_parsed", None

    def add_materials(self, project_ref: str, material_paths: Iterable[str], added_via: str) -> list[dict]:
        project_record = self.get_project_record(project_ref)
        if not project_record:
            raise ValueError(f"项目 {project_ref} 不存在")

        project_path = Path(project_record["project_dir"]).resolve()
        workspace_root = Path(project_record["workspace_dir"]).resolve()
        materials = self._load_materials(project_record)
        added_materials: list[dict] = []

        for raw_path in material_paths:
            source_path = Path(raw_path).expanduser().resolve()
            if not source_path.exists() or not source_path.is_file():
                raise ValueError(f"材料不存在: {raw_path}")

            # Fix4: hard size limit applies to BOTH paths. Live workspace-selected files used to
            # skip this (the check sat only in the import/copy branch), letting an oversized
            # in-workspace file slip past the 25MB cap. Enforce BEFORE the workspace-vs-import
            # branch so neither path can reference/copy an oversized source.
            from backend import material_limits as _ml
            source_size = source_path.stat().st_size
            if source_size > _ml.MAX_HEAVY_MATERIAL_BYTES:
                limit_mb = _ml.MAX_HEAVY_MATERIAL_BYTES / (1024 * 1024)
                actual_mb = source_size / (1024 * 1024)
                raise ValueError(
                    f"文件 {source_path.name!r} 大小 {actual_mb:.1f} MB 超过上传限制 "
                    f"{limit_mb:.0f} MB，请压缩后重试"
                )

            workspace_relative = self._workspace_relative_path(source_path, workspace_root)
            if workspace_relative is not None:
                duplicate = self._find_existing_workspace_material(materials, workspace_relative)
                if duplicate:
                    added_materials.append(duplicate)
                    continue
                stored_rel_path = workspace_relative
                source_type = "workspace"
                original_path = ""
                # workspace 材料是 live 文件，按源路径取 add-time hash（已知低危：之后被改写不重算，v1 接受）
                hashed_path = source_path
            else:
                duplicate = self._find_existing_imported_material(materials, source_path)
                if duplicate:
                    added_materials.append(duplicate)
                    continue
                destination_rel = self._build_imported_destination(project_path, source_path.name)
                destination_path = project_path / destination_rel
                destination_path.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(source_path, destination_path)
                stored_rel_path = self._to_posix(destination_rel)
                source_type = "imported"
                original_path = str(source_path)
                # chat_upload（imported）已落盘拷贝，按拷贝后文件取 hash
                hashed_path = destination_path

            mime_type, _ = mimetypes.guess_type(source_path.name)
            material = {
                "id": self._new_id("mat"),
                "display_name": source_path.name,
                "media_kind": self._detect_media_kind(source_path),
                "source_type": source_type,
                "stored_rel_path": stored_rel_path,
                "original_path": original_path,
                "added_via": added_via,
                "file_type": source_path.suffix.lstrip(".").lower(),
                "mime_type": mime_type or "application/octet-stream",
                "size_bytes": (source_path.stat().st_size if source_path.exists() else 0),
                "content_sha256": self._content_sha256(hashed_path),
                "created_at": datetime.now().isoformat(timespec="seconds"),
            }
            materials.append(material)
            added_materials.append(material)

        self._save_materials(project_record, materials)
        self._touch_project(project_record["id"])
        return added_materials

    def remove_material(self, project_ref: str, material_id: str):
        project_record = self.get_project_record(project_ref)
        if not project_record:
            raise ValueError(f"项目 {project_ref} 不存在")

        materials = self._load_materials(project_record)
        target = next((item for item in materials if item["id"] == material_id), None)
        if not target:
            raise ValueError("材料不存在")

        # 删源文件前先 release：shared-hash 缓存仅在最后一个引用消失时才真删。
        # 关键：用 live-file 内容算 key（与 retain 一致），且必须在源文件还在时算，否则 hash 不上。
        converter = getattr(self, "_material_converter", None)
        if converter is not None:
            target_path = self.get_material_path(project_ref, target["id"])
            if target_path.exists():
                converter.release(self._cache_key_for_material(target, target_path), target["id"])

        if target["source_type"] == "imported":
            imported_path = Path(project_record["project_dir"]) / target["stored_rel_path"]
            if imported_path.exists():
                imported_path.unlink()

        materials = [item for item in materials if item["id"] != material_id]
        self._save_materials(project_record, materials)
        self._touch_project(project_record["id"])

    def delete_project(self, project_ref: str):
        project_record = self.get_project_record(project_ref)
        if project_record:
            self._release_project_material_caches(project_record)   # N6: free cache refs before deleting the project dir
            project_path = Path(project_record["project_dir"])
            if project_path.exists():
                shutil.rmtree(project_path)
            registry = self._load_registry()
            registry["projects"] = [item for item in registry["projects"] if item["id"] != project_record["id"]]
            self._save_registry(registry)
            return


        raise ValueError(f"项目 {project_ref} 不存在")

    def _release_project_material_caches(self, project_record: dict) -> None:
        """N6: 删项目前，逐条释放材料的共享缓存引用——N6 缓存活在 projects_dir 之外
        （materials_cache），rmtree 删不到。release 必须在源文件还在时执行（与 remove_material 同），
        最后一个引用消失时 shared-hash GC 才会删掉转写/markdown/tombstone/.refs。
        无 converter / 缺路径优雅跳过，单条失败不阻断删项目。"""
        converter = getattr(self, "_material_converter", None)
        if converter is None:
            return
        try:
            materials = self._load_materials(project_record)
        except Exception:
            return
        for material in materials:
            try:
                if not material.get("content_sha256"):
                    continue
                path = self._resolve_material_path(project_record, material)
                if not path.exists():
                    continue
                converter.release(self._cache_key_for_material(material, path), material["id"])
            except Exception:  # noqa: BLE001 单个材料 release 失败不应阻断删项目
                continue

    def read_file(self, project_ref: str, file_path: str) -> str:
        """璇诲彇椤圭洰鏂囦欢"""
        project_path = self.get_project_path(project_ref)
        if not project_path:
            raise ValueError(f"项目 {project_ref} 不存在")

        full_path = self._resolve_project_path(project_path, file_path)
        if not full_path.exists():
            raise ValueError(f"文件 {file_path} 不存在")

        return full_path.read_text(encoding="utf-8")

    def write_file(self, project_ref: str, file_path: str, content: str):
        """写入项目文件（原子：同目录 temp + os.replace，避免并发读到写入中间态）"""
        project_path = self.get_project_path(project_ref)
        if not project_path:
            raise ValueError(f"项目 {project_ref} 不存在")

        normalized_path = self.validate_plan_write(project_ref, file_path)
        full_path = self._resolve_project_path(project_path.resolve(), normalized_path)
        full_path.parent.mkdir(parents=True, exist_ok=True)
        tmp_fd, tmp_name = tempfile.mkstemp(dir=str(full_path.parent), suffix=".tmp")
        os.close(tmp_fd)
        try:
            Path(tmp_name).write_text(content, encoding="utf-8")  # 与原 write_text 同 newline 行为
            os.replace(tmp_name, full_path)
        except BaseException:
            try:
                os.unlink(tmp_name)
            except OSError:
                pass
            raise

    def list_workspace_files(self, project_ref: str) -> list[dict]:
        """R3: structured workspace file list for the front-end file tree.
        Skips retired files and everything under materials/. Each .md → {path, group,
        stage, editable, mtime_ns}. mtime_ns is a str (opaque — never coerce to Number;
        JS loses precision past 2^53)."""
        project_path = self.get_project_path(project_ref)
        if not project_path:
            raise ValueError(f"项目 {project_ref} 不存在")

        files = []
        for md_file in project_path.rglob("*.md"):
            rel_path = self._to_posix(md_file.relative_to(project_path)).lstrip("/")
            if rel_path in self.RETIRED_WORKSPACE_FILES:
                continue
            if rel_path.startswith("materials/"):
                continue
            try:
                mtime_ns = str(md_file.stat().st_mtime_ns)
            except OSError:
                # 文件在 rglob 枚举后、stat 前被并发删除/改名（AI 改写期间）——跳过，
                # 不让整个列表 500；刷新自愈。
                continue
            semantics = self.get_file_semantics(rel_path)
            files.append({
                "path": rel_path,
                "group": semantics["group"],
                "stage": semantics["stage"],
                "editable": semantics["editable"],
                "mtime_ns": mtime_ns,
            })
        return files

    def read_file_with_mtime(self, project_ref: str, file_path: str) -> dict:
        """R3: content + mtime_ns for the edit base. NO lock here (the per-project request
        lock is held by chat_stream for a full turn — locking reads would freeze preview).
        stat BEFORE read so the returned base_mtime is never NEWER than the bytes returned:
        if an AI write interleaves, the worst case is a safe 409 on save (user reloads),
        never a silent overwrite. The AI writes EDITABLE files only via the atomic write_file
        (os.replace), so a no-lock read of an editable file never sees a half-written one.
        (Read-only tracking files are still direct-written; a rare torn preview self-heals on
        reload — they are never editable nor in the save/CAS path.)"""
        project_path = self.get_project_path(project_ref)
        if not project_path:
            raise ValueError(f"项目 {project_ref} 不存在")
        full_path = self._resolve_project_path(project_path, file_path)
        if not full_path.exists():
            raise ValueError(f"文件 {file_path} 不存在")
        mtime_ns = str(full_path.stat().st_mtime_ns)
        content = full_path.read_text(encoding="utf-8")
        return {"content": content, "mtime_ns": mtime_ns}

    def user_write_file(self, project_ref: str, file_path: str, content: str,
                        base_mtime_ns: str) -> str:
        """R3: atomic user write with mtime CAS. Caller MUST hold the per-project request
        lock (shared with chat writes) so the stat→replace window is not racing an AI write.
        Returns new mtime_ns (str). Raises:
          - ValueError('非法的文件路径') on traversal           → endpoint 400
          - UserWriteForbiddenError on non-whitelisted file     → endpoint 403
          - FileNotFoundError on missing file                   → endpoint 404
          - StaleFileError on mtime mismatch                    → endpoint 409
          - OSError on the atomic write itself (target locked)  → endpoint 500
        """
        project_path = self.get_project_path(project_ref)
        if not project_path:
            raise ValueError(f"项目 {project_ref} 不存在")
        canonical = self.validate_user_write(project_ref, file_path)  # UserWriteForbiddenError / ValueError
        full_path = self._resolve_project_path(project_path.resolve(), canonical)
        if not full_path.exists():
            raise FileNotFoundError(f"文件 {canonical} 不存在")
        current_mtime_ns = str(full_path.stat().st_mtime_ns)
        if current_mtime_ns != base_mtime_ns:
            raise StaleFileError(current_mtime_ns)
        # 原子写：同目录 temp + os.replace（与 write_file 同款，newline 行为一致）
        tmp_fd, tmp_name = tempfile.mkstemp(dir=str(full_path.parent), suffix=".tmp")
        os.close(tmp_fd)
        try:
            Path(tmp_name).write_text(content, encoding="utf-8")
            os.replace(tmp_name, full_path)
        except BaseException:
            try:
                os.unlink(tmp_name)
            except OSError:
                pass
            raise
        return str(full_path.stat().st_mtime_ns)

    def normalize_file_path(self, project_ref: str, file_path: str) -> str:
        project_path = self.get_project_path(project_ref)
        if not project_path:
            raise ValueError(f"项目 {project_ref} 不存在")

        project_root = project_path.resolve()
        full_path = self._resolve_project_path(project_root, file_path)
        normalized_path = self._to_posix(full_path.relative_to(project_root))
        return self._canonicalize_plan_markdown_path(normalized_path)

    def is_formal_plan_file(self, file_path: str) -> bool:
        normalized_path = self._canonicalize_plan_markdown_path(self._to_posix(file_path).lstrip("/"))
        if not self._is_plan_markdown_path(normalized_path):
            return False
        return normalized_path.split("/", 1)[1] in self.FORMAL_PLAN_FILES

    def evidence_gate_satisfied(self, project_ref: str) -> bool:
        project_path = self.get_project_path(project_ref)
        if not project_path:
            raise ValueError(f"项目 {project_ref} 不存在")
        return self._has_effective_notes(project_path) and self._has_effective_references(project_path)

    def validate_plan_write(self, project_ref: str, file_path: str) -> str:
        normalized_path = self.normalize_file_path(project_ref, file_path)

        if not self._is_plan_markdown_path(normalized_path):
            return normalized_path

        if self._is_backend_owned_stage_tracking_file(normalized_path):
            raise ValueError(
                f"`{normalized_path}` is backend-generated. Update the substantive project files instead of writing stage tracking files directly."
            )

        if not self.is_formal_plan_file(normalized_path):
            raise ValueError(
                f"`{normalized_path}` is not an official plan file. Use only the registered `plan/*.md` files and never invent unofficial files such as `plan/gate-control.md`."
            )

        if self._requires_pre_outline_evidence(normalized_path) and not self.evidence_gate_satisfied(project_ref):
            raise ValueError(
                "Before writing `plan/outline.md` or `plan/research-plan.md`, update `notes.md` and `references.md` and satisfy the minimum 2-source rule."
            )

        return normalized_path

    def _canonical_user_path(self, normalized_path: str) -> str:
        # 注意：不复用 _canonicalize_plan_markdown_path（它只 lower plan/*.md，content/ 不动）。
        # 这里对整条 posix 相对路径统一 casefold——Windows 文件系统大小写不敏感，
        # content/Report_Draft_V1.MD 必须与 content/report_draft_v1.md 判为同一文件。
        return self._to_posix(normalized_path).lstrip("/").casefold()

    def is_user_editable(self, normalized_path: str) -> bool:
        return self._canonical_user_path(normalized_path) in self.USER_EDITABLE_FILES

    def get_file_semantics(self, normalized_path: str) -> dict:
        """Map a normalized relative path to {group, stage, editable}.
        Unknown .md → group='other', stage=None, editable=False."""
        canonical = self._canonical_user_path(normalized_path)
        semantics = self.FILE_SEMANTICS.get(canonical, {"group": "other", "stage": None})
        return {
            "group": semantics["group"],
            "stage": semantics["stage"],
            "editable": canonical in self.USER_EDITABLE_FILES,
        }

    def validate_user_write(self, project_ref: str, file_path: str) -> str:
        """R3: independent whitelist gate for USER (HTTP) writes — NOT validate_plan_write
        (that carries the LLM-only pre-outline evidence gate and does not itself deny
        independent-review/lint-report; those live in the chat tool layer the HTTP endpoint
        never reaches). Whitelist = default-deny.
        Path traversal → ValueError (endpoint 400). Not whitelisted → UserWriteForbiddenError
        (403) — a distinct type, NOT PermissionError, so a filesystem PermissionError from the
        actual write is not misread as a whitelist denial.
        Returns the whitelist canonical path so the write target is stable across casing."""
        normalized = self.normalize_file_path(project_ref, file_path)  # 穿越路径在此抛 ValueError
        canonical = self._canonical_user_path(normalized)
        if canonical not in self.USER_EDITABLE_FILES:
            raise UserWriteForbiddenError(f"`{normalized}` 不可由用户手动编辑")
        return canonical

    def _delivery_log_has_placeholder_feedback(self, content: str) -> bool:
        if self._DELIVERY_PLACEHOLDER_INLINE.search(content):
            return True
        for match in self._DELIVERY_BLOCK_RE.finditer(content):
            body = match.group("body")
            for line in body.splitlines():
                stripped = line.strip()
                if stripped.startswith("- [") or stripped.startswith("#"):
                    break
                if self._PLACEHOLDER_WORDS_RE.search(line):
                    return True
        return False

    def validate_self_signature(self, normalized_path: str, content: str, checkpoints: dict) -> str | None:
        """Return an error message if `content` violates self-signature / premature-verdict /
        archive-claim rules for the given plan file path, else None.

        Interception is auto-disabled when the corresponding checkpoint stamp is present
        (spec §12.3): `review_passed_at` disables review-checklist.md interception; and
        `delivery_archived_at` disables delivery-log.md interception. The auto-disable
        lets the UI-driven advance flow write whatever it needs without false positives
        after the user has already confirmed the stage via the right-side workspace.

        Returns:
            str error message for the caller to surface in the chat stream AND return
            as the tool error, or None if the write is allowed.
        """
        if normalized_path == "plan/review-checklist.md":
            if "review_passed_at" in checkpoints:
                return None
            for pattern in self._SELF_SIGNATURE_PATTERNS:
                if pattern.search(content):
                    return (
                        "review-checklist.md 的\"审查人\"字段必须由真实用户签字，"
                        "请保留\"审查人：[待用户确认]\"让用户在 UI 上签字。"
                    )
            if "review_started_at" not in checkpoints:
                for pattern in self._PREMATURE_REVIEW_VERDICT_PATTERNS:
                    if pattern.search(content):
                        return (
                            "review-checklist.md 的\"审查结论 / 建议通过\"字段必须在用户点击"
                            "\"完成撰写，开始审查\"按钮之后再写入。当前审查尚未开始，"
                            "请保留为空或\"[待审查]\"，并告知用户需要他们先点按钮进入审查阶段。"
                        )
        if normalized_path == "plan/delivery-log.md":
            if "delivery_archived_at" in checkpoints:
                return None
            for pattern in self._ARCHIVE_CLAIM_PATTERNS:
                if pattern.search(content):
                    return (
                        "delivery-log.md 声明\"已归档/已交付\"需要用户点击 UI 的\"归档结束项目\"按钮。"
                        "请把状态保持为\"待归档\"，并告知用户需要他们点按钮。"
                    )
            if self._delivery_log_has_placeholder_feedback(content):
                return (
                    "delivery-log.md 勾选\"客户反馈\"需要真实反馈内容，"
                    "请保留为未勾选，等用户补齐反馈后再勾。"
                )
        return None

    def is_protected_stage_checkpoints_path(self, normalized_path: str) -> bool:
        """Return True if `normalized_path` points to `stage_checkpoints.json` (the user-
        confirmation truth source). The model must never be able to directly write this
        file via `write_file` — only the checkpoint endpoints (via `record_stage_checkpoint`)
        may mutate it.

        Comparison must be case-insensitive because Windows filesystems are case-insensitive
        by default: `Stage_Checkpoints.json` and `STAGE_CHECKPOINTS.JSON` all resolve to
        the same file. Backslashes are normalized to forward slashes before the basename
        extraction so Windows-style relative paths (`plan\\..\\Stage_Checkpoints.json`)
        are also blocked.
        """
        if not normalized_path:
            return False
        tail = normalized_path.replace("\\", "/").rsplit("/", 1)[-1]
        return tail.casefold() == self.STAGE_CHECKPOINTS_FILENAME.casefold()

    def read_material_file(self, project_ref: str, material_id: str) -> str:
        from backend import material_limits

        project_record = self.get_project_record(project_ref)
        if not project_record:
            raise ValueError(f"项目 {project_ref} 不存在")

        material = self.get_material(project_ref, material_id)
        material_path = self.get_material_path(project_ref, material_id)
        suffix = material_path.suffix.lower()

        if material_limits.is_heavy_suffix(suffix):
            actual = material_path.stat().st_size if material_path.exists() else 0
            if actual > material_limits.MAX_HEAVY_MATERIAL_BYTES:
                raise ValueError(
                    "这个文件过大，读不动；请只传关键的评分标准/技术规范书等小文件"
                )

        if material["media_kind"] == "image_like":
            text = self._converter_read_image(project_ref, material_id)
        else:
            text = self._converter_read_document(project_ref, material_path)

        self._retain_material_cache(material, material_path)
        return text

    def set_material_converter(self, converter):
        self._material_converter = converter

    @staticmethod
    def _content_sha256(path: Path) -> str:
        """落盘文件 sha256（与 MaterialConverter._content_hash 同算法、同分块）。"""
        import hashlib

        h = hashlib.sha256()
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(1 << 20), b""):
                h.update(chunk)
        return h.hexdigest()

    def _cache_key_for_material(self, material: dict, material_path: Path) -> str:
        """用当前 live 文件内容算 converter 缓存 key（与 transcribe_image/convert_document 的 live-hash 一致），
        避免 add-time content_sha256 与改动后的 workspace 文件分歧。
        文档 extra=""、图片 extra=converter.image_cache_extra。"""
        extra = (
            self._material_converter.image_cache_extra
            if material.get("media_kind") == "image_like"
            else ""
        )
        content_hash = self._content_sha256(material_path)
        return self._material_converter.cache_key_from_sha256(content_hash, extra)

    def _retain_material_cache(self, material: dict, material_path: Path) -> None:
        converter = getattr(self, "_material_converter", None)
        if converter is None or not material_path.exists():
            return
        converter.retain(self._cache_key_for_material(material, material_path), material["id"])

    def retain_material_cache(self, project_ref, material_id) -> None:
        """chat 路径专用：当前轮自己 transcribe 图片后补 retain（live-hash key），
        否则同内容另一材料被删时会连带删掉共享缓存。retain 失败由调用方吞掉、不影响展示。"""
        converter = getattr(self, "_material_converter", None)
        if converter is None:
            return
        try:
            material = self.get_material(project_ref, material_id)
            material_path = self.get_material_path(project_ref, material_id)
        except Exception:
            return
        if not material_path.exists():
            return
        converter.retain(self._cache_key_for_material(material, material_path), material["id"])

    def _converter_read_document(self, project_ref, material_path):
        if getattr(self, "_material_converter", None) is None:
            return self._legacy_read_document(material_path)   # 无 converter 的纯单测回退
        from backend.material_conversion import MaterialConversionError

        try:
            return self._material_converter.convert_document(material_path)
        except MaterialConversionError as exc:
            raise ValueError(str(exc)) from exc

    def _converter_read_image(self, project_ref, material_id):
        converter = getattr(self, "_material_converter", None)
        if converter is None:
            raise ValueError("当前环境无法读取图片材料")
        from backend.material_conversion import MaterialConversionError

        material_path = self.get_material_path(project_ref, material_id)
        material = self.get_material(project_ref, material_id)
        mime = material.get("mime_type") or "image/png"
        try:
            return converter.transcribe_image(material_path, mime)
        except MaterialConversionError as exc:
            raise ValueError(str(exc)) from exc

    def _legacy_read_document(self, material_path: Path) -> str:
        suffix = material_path.suffix.lower()

        if suffix in self.TEXT_SUFFIXES:
            return material_path.read_text(encoding="utf-8")
        if suffix == ".docx":
            return self._read_docx(material_path)
        if suffix == ".xlsx":
            return self._read_xlsx(material_path)
        if suffix == ".pdf":
            return self._read_pdf(material_path)

        try:
            return material_path.read_text(encoding="utf-8")
        except UnicodeDecodeError as exc:
            raise ValueError(f"当前暂不支持读取 {suffix} 材料") from exc

    def _get_raw_material(self, project_record: dict, material_id: str) -> dict:
        """N6 Fix4（perf）：按 id 取 RAW 材料记录，不做 conversion_status 富化。
        chat 路径 / read_material_file 只需路径与元字段，不该为单条查找付 list_materials 的
        O(n) 状态探测。status 富化只在公开 list_materials（/materials 端点）里发生。"""
        material = next(
            (item for item in self._load_materials(project_record) if item["id"] == material_id),
            None,
        )
        if not material:
            raise ValueError("材料不存在")
        return material

    def get_material(self, project_ref: str, material_id: str) -> dict:
        project_record = self.get_project_record(project_ref)
        if not project_record:
            raise ValueError(f"项目 {project_ref} 不存在")
        return self._get_raw_material(project_record, material_id)

    def get_material_path(self, project_ref: str, material_id: str) -> Path:
        project_record = self.get_project_record(project_ref)
        if not project_record:
            raise ValueError(f"项目 {project_ref} 不存在")

        material = self._get_raw_material(project_record, material_id)
        return self._resolve_material_path(project_record, material)

    @staticmethod
    def _resolve_material_path(project_record: dict, material: dict) -> Path:
        """从 material dict + project_record 直接解析落盘路径，不经 get_material（避免 list_materials 再入）。"""
        if material["source_type"] == "workspace":
            return (Path(project_record["workspace_dir"]) / material["stored_rel_path"]).resolve()
        return (Path(project_record["project_dir"]) / material["stored_rel_path"]).resolve()

    def get_workspace_summary(self, project_ref: str) -> dict:
        project_path = self.get_project_path(project_ref)
        if not project_path:
            raise ValueError(f"项目 {project_ref} 不存在")

        self._backfill_stage_checkpoints_if_missing(project_path)
        stage_state = self._infer_stage_state(project_path)
        project_record = self.get_project_record(project_ref) or {}
        tracking_state = self._sync_stage_tracking_files(project_path)
        materials = self.list_materials(project_ref)

        checkpoints = stage_state.get("checkpoints", {})
        next_stage_hint = None
        if "review_passed_at" in checkpoints:
            next_stage_hint = "S6" if self._delivery_mode_requires_presentation(project_path) else "S7"

        stalled_since = None
        if stage_state["stage_code"] in ("S2", "S3"):
            last_write = self._last_evidence_write_at(project_path)
            if last_write is not None:
                elapsed = datetime.now() - last_write
                if elapsed.total_seconds() >= 30 * 60:
                    stalled_since = last_write.isoformat(timespec="seconds")

        length_targets = stage_state.get("length_targets", {})

        return {
            "stage_code": stage_state["stage_code"],
            "status": stage_state.get("stage_status", tracking_state["status"]),
            "completed_items": stage_state["completed_items"],
            "skipped_items": stage_state.get("skipped_items", []),
            "next_actions": tracking_state["next_actions"],
            "workspace_dir": project_record.get("workspace_dir", ""),
            "project_dir": str(project_path),
            "materials": materials,
            "checkpoints": checkpoints,
            "length_targets": length_targets,
            "length_fallback_used": length_targets.get("fallback_used", False),
            "quality_progress": self._build_quality_progress(project_path, stage_state),
            "flags": {
                **stage_state.get("flags", {}),
                "review_stale": self._is_report_review_stale(project_path),
            },
            "next_stage_hint": next_stage_hint,
            "stalled_since": stalled_since,
            "word_count": self._current_report_word_count(project_path),
            "delivery_mode": self._extract_delivery_mode(project_path),
        }

    def build_project_context(self, project_ref: str) -> str:
        project_path = self.get_project_path(project_ref)
        if not project_path:
            raise ValueError(f"项目 {project_ref} 不存在")

        self._sync_stage_tracking_files(project_path)
        sections = []
        for title, relative_path in self.CORE_CONTEXT_FILES:
            content = self._read_optional(project_ref, relative_path)
            if content:
                sections.append(f"## {title}\n{content}")

        if self._is_effective_plan_file(project_path, "tasks.md"):
            tasks_content = self._read_plan_file(project_path, "tasks.md")
            if tasks_content:
                sections.append(f"## 当前阶段任务\n{tasks_content}")

        materials = self.list_materials(project_ref)
        if materials:
            material_lines = [
                f"- {material['id']} | {material['display_name']} | {material['source_type']} | {material['file_type']}"
                for material in materials
            ]
            sections.append("## 可用项目材料\n" + "\n".join(material_lines))

        return "\n\n".join(sections)

    def get_script_path(self, script_name: str) -> str:
        script_path = self.skill_dir / "scripts" / script_name
        if not script_path.exists():
            raise ValueError(f"脚本 {script_name} 不存在")
        return str(script_path)

    def ensure_output_dir(self, project_ref: str) -> str:
        project_path = self.get_project_path(project_ref)
        if not project_path:
            raise ValueError(f"项目 {project_ref} 不存在")

        output_dir = project_path / "output"
        output_dir.mkdir(parents=True, exist_ok=True)
        return str(output_dir)

    def get_primary_report_path(self, project_ref: str) -> str:
        project_path = self.get_project_path(project_ref)
        if not project_path:
            raise ValueError(f"项目 {project_ref} 不存在")

        report_path = project_path / self.REPORT_DRAFT_PATH
        if report_path.exists():
            return str(report_path)

        raise ValueError(f"没有可检查或导出的报告草稿，请先生成 {self.REPORT_DRAFT_PATH}")

    def _ensure_stage_gates_state(self, project_path: Path) -> str:
        return self._sync_stage_tracking_files(project_path)["stage_gates_text"]
        stage_gates_path = project_path / "plan" / "stage-gates.md"
        if not stage_gates_path.exists():
            template_path = self.skill_dir / "plan-template" / "stage-gates.md"
            if template_path.exists():
                shutil.copy(template_path, stage_gates_path)
            else:
                return ""

        stage_state = self._infer_stage_state(project_path)
        stage_code = stage_state["stage_code"]
        if not stage_code:
            return stage_gates_path.read_text(encoding="utf-8")

        original_content = stage_gates_path.read_text(encoding="utf-8")
        content = original_content
        content = re.sub(r"\*\*闃舵\*\*:\s*[^\n]+", f"**闃舵**: {stage_code}", content)
        content = re.sub(r"\*\*状态\*\*:\s*[^\n]+", "**状态**: 进行中", content)
        content = re.sub(
            r"\*\*鏇存柊鏃ユ湡\*\*:\s*[^\n]+",
            f"**鏇存柊鏃ユ湡**: {datetime.now().strftime('%Y-%m-%d')}",
            content,
        )
        for task in self._tracked_stage_items():
            content = self._set_stage_gate_item_state(content, task, " ")
        for task in stage_state["completed_items"]:
            content = self._set_stage_gate_item_state(content, task, "x")
        for task in stage_state["skipped_items"]:
            content = self._set_stage_gate_item_state(content, task, "/")

        if content != original_content:
            stage_gates_path.write_text(content, encoding="utf-8")
        return content

    def _sync_stage_tracking_files(self, project_path: Path) -> dict:
        stage_state = self._infer_stage_state(project_path)
        stage_code = stage_state["stage_code"] or "S0"
        completed_items = list(stage_state["completed_items"])
        skipped_items = list(stage_state["skipped_items"])
        next_actions = [
            item
            for item in self.STAGE_CHECKLIST_ITEMS.get(stage_code, [])
            if item not in completed_items and item not in skipped_items
        ]
        if stage_code == "S5":
            flags = stage_state.get("flags", {})
            independent_ready = bool(flags.get("independent_review_ready"))
            lint_ready = bool(flags.get("lint_report_ready"))
            if not independent_ready and not lint_ready:
                next_actions = ["请点击上方'独立审查'和'AI 味自查'按钮"]
            elif independent_ready and not lint_ready:
                next_actions = ["还差'AI 味自查'，请点击上方按钮"]
            elif lint_ready and not independent_ready:
                next_actions = ["还差'独立审查'，请点击上方按钮"]
            else:
                next_actions = ["等主代理跟你讨论审查结果，确认通过后说'审查通过'"]
        status = "进行中"
        stage_gates_path = project_path / "plan" / "stage-gates.md"
        existing_stage_gates = stage_gates_path.read_text(encoding="utf-8") if stage_gates_path.exists() else ""
        manual_lines = self._extract_manual_stage_gate_lines(existing_stage_gates)
        stage_gates_text = self._render_stage_gates_markdown(
            stage_code,
            status,
            completed_items,
            skipped_items,
            manual_lines,
        )
        progress_text = self._render_progress_markdown(
            stage_code, status, next_actions, completed_items,
            stage_state=stage_state,
        )
        tasks_text = self._render_tasks_markdown(stage_code, next_actions)

        self._write_tracking_file(stage_gates_path, stage_gates_text)
        self._write_tracking_file(project_path / "plan" / "progress.md", progress_text)
        self._write_tracking_file(project_path / "plan" / "tasks.md", tasks_text)

        return {
            "stage_code": stage_code,
            "status": status,
            "completed_items": completed_items,
            "skipped_items": skipped_items,
            "next_actions": next_actions,
            "stage_gates_text": stage_gates_text,
        }

    def _extract_delivery_mode(self, project_path: Path) -> str:
        """Parse 交付形式 from plan/project-overview.md."""
        overview_path = project_path / "plan" / "project-overview.md"
        if not overview_path.exists():
            return "仅报告"
        text = overview_path.read_text(encoding="utf-8")
        match = re.search(r"交付形式[^\n]*?[:：]\s*([^\n]+)", text)
        if not match:
            return "仅报告"
        value = match.group(1).strip()
        return "报告+演示" if "演示" in value else "仅报告"

    def _current_report_word_count(self, project_path: Path) -> int:
        draft_text = self._read_project_file(project_path, self.REPORT_DRAFT_PATH)
        if not draft_text or self._is_template_stub(draft_text):
            return 0
        return self._count_words(draft_text)

    def _last_evidence_write_at(self, project_path: Path) -> datetime | None:
        candidates = [
            project_path / "plan" / "notes.md",
            project_path / "plan" / "references.md",
            project_path / "plan" / "data-log.md",
            project_path / "plan" / "analysis-notes.md",
        ]
        mtimes = [path.stat().st_mtime for path in candidates if path.exists()]
        if not mtimes:
            return None
        return datetime.fromtimestamp(max(mtimes))

    def _build_quality_progress(self, project_path: Path, stage_state: dict) -> dict | None:
        stage = stage_state["stage_code"]
        targets = stage_state.get("length_targets", {})
        if stage == "S2":
            return {
                "label": "有效来源条目",
                "current": self._count_valid_data_log_sources(project_path),
                "target": targets.get("data_log_min", 0),
            }
        if stage == "S3":
            return {
                "label": "分析证据引用",
                "current": self._count_analysis_refs(project_path),
                "target": targets.get("analysis_refs_min", 0),
            }
        return None

    def record_stage_checkpoint(self, project_id: str, key: str, action: str) -> dict:
        from backend.chat import _get_project_request_lock

        project_path = self.get_project_path(project_id)
        if project_path is None:
            raise ValueError(f"项目不存在: {project_id}")
        if action not in ("set", "clear"):
            raise ValueError(f"未知 action: {action}")
        lock = _get_project_request_lock(project_id)
        with lock:
            if key == "review_passed_at" and action == "set":
                from backend.independent_review import get_independent_review_lock
                from backend.report_tools import get_lint_report_lock

                review_lock = get_independent_review_lock(project_id)
                if review_lock.locked():
                    raise ValueError("独立审查正在进行中，请等待完成后再标记审查通过")

                lint_lock = get_lint_report_lock(project_id)
                if lint_lock.locked():
                    raise ValueError("AI 味自查正在进行中，请等待完成后再标记审查通过")

            if action == "set":
                self._validate_stage_checkpoint_transition(project_path, key)
                timestamp = self._save_stage_checkpoint(project_path, key)
                # R5: 确认大纲时若当前无有效方法论快照（首次确认 / 上次快照写入失败留下的半提交）
                # → 写/补快照；已有有效快照则不重写（防确认后改 outline 声明行静默换方法论，红队
                # BLOCKER 2）。用「当前快照状态」而非「是否首次确认」判定，使快照写入失败后下次确认
                # 可自愈补写，消除两阶段写的永久半提交不一致（codex B4 红队 BLOCKER）。
                if key == "outline_confirmed_at":
                    snap_state, _ = self.read_confirmed_methodology_snapshot(project_path)
                    if snap_state != "parsed":
                        self._snapshot_methodology_on_confirm(project_path)
                self._sync_stage_tracking_files(project_path)
                return {"status": "ok", "key": key, "timestamp": timestamp}
            self._clear_stage_checkpoint_cascade(project_path, key)
            self._sync_stage_tracking_files(project_path)
            return {"status": "ok", "key": key, "cleared": True}

    def _write_tracking_file(self, path: Path, content: str) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        existing = path.read_text(encoding="utf-8") if path.exists() else ""
        if existing != content:
            path.write_text(content, encoding="utf-8")

    def _extract_manual_stage_gate_lines(self, markdown_text: str) -> list[str]:
        manual_lines: list[str] = []
        tracked_items = set(self._tracked_stage_items())
        for raw_line in markdown_text.splitlines():
            stripped = raw_line.strip()
            match = re.match(r"^- \[(?: |x|X|/)\] (.+)$", stripped)
            if not match:
                continue
            task = match.group(1).strip()
            if task in tracked_items:
                continue
            manual_lines.append(stripped)
        return manual_lines

    def _render_stage_gates_markdown(
        self,
        stage_code: str,
        status: str,
        completed_items: list[str],
        skipped_items: list[str],
        manual_lines: list[str] | None = None,
    ) -> str:
        lines = [
            "# 项目阶段与门禁",
            "",
            "## 当前阶段",
            "",
            f"**阶段**: {stage_code}",
            f"**状态**: {status}",
            f"**更新日期**: {datetime.now().strftime('%Y-%m-%d')}",
            "",
            "## 阶段进度",
            "",
        ]
        for stage in self.STAGE_ORDER:
            lines.append(f"### {stage} {self.STAGE_TITLES.get(stage, stage)}")
            for item in self.STAGE_CHECKLIST_ITEMS[stage]:
                state = "x" if item in completed_items else "/" if item in skipped_items else " "
                lines.append(f"- [{state}] {item}")
            lines.append("")
        if manual_lines:
            lines.append("## 手工补充记录")
            lines.append("")
            lines.extend(manual_lines)
            lines.append("")
        return "\n".join(lines).strip() + "\n"

    def _render_progress_markdown(
        self,
        stage_code: str,
        status: str,
        next_actions: list[str],
        completed_items: list[str],
        *,
        stage_state: dict | None = None,
    ) -> str:
        current_task = next_actions[0] if next_actions else "当前阶段任务已完成，等待推进下一阶段。"
        completed_summary = " / ".join(completed_items[-3:]) if completed_items else "-"
        next_summary = " / ".join(next_actions[:3]) if next_actions else "-"
        lines = [
            "# 项目进度追踪",
            "",
            "## 当前状态",
            f"**阶段**: {stage_code}",
            f"**状态**: {status}",
            f"**当前任务**: {current_task}",
        ]
        # v5: S2/S3 阶段渲染 quality_progress 行
        if stage_state and stage_code in {"S2", "S3"}:
            qp = stage_state.get("quality_progress")
            if qp and isinstance(qp.get("target"), int) and qp["target"] > 0:
                label = qp.get("label", "")
                current = qp.get("current", 0)
                target = qp["target"]
                lines.append(f"**质量进度**: {current}/{target} {label}")
        lines.extend([
            f"**更新日期**: {datetime.now().strftime('%Y-%m-%d')}",
            "",
            "## 执行摘要",
            f"- 已完成: {completed_summary}",
            f"- 下一步: {next_summary}",
        ])
        return "\n".join(lines).strip() + "\n"

    def _render_tasks_markdown(self, stage_code: str, next_actions: list[str]) -> str:
        lines = [
            "# 任务清单",
            "",
            "## 当前阶段",
            "",
            f"**阶段**: {stage_code}",
            f"**阶段目标**: {self.STAGE_TITLES.get(stage_code, stage_code)}",
            "",
            "## 当前阶段待办",
        ]
        if next_actions:
            lines.extend(f"- [ ] {item}" for item in next_actions)
        else:
            lines.append("- [x] 当前阶段待办已清空，可推进下一阶段。")
        return "\n".join(lines).strip() + "\n"

    def _infer_stage_progress(self, project_path: Path) -> tuple[str, list[str]]:
        stage_state = self._infer_stage_state(project_path)
        return stage_state["stage_code"], stage_state["completed_items"]

    def _has_meaningful_outline(self, project_path: Path) -> bool:
        return self._has_effective_outline(project_path)

    def _infer_stage_state(self, project_path: Path) -> dict:
        targets = self._resolve_length_targets(project_path)
        checkpoints = self._load_stage_checkpoints(project_path)
        stage_one_state = self._stage_one_completion_state(project_path, checkpoints)
        stage_four_state = self._stage_four_completion_state(project_path, checkpoints, targets, stage_one_state)
        stage_five_state = self._stage_five_completion_state(
            project_path, checkpoints, targets, stage_one_state, stage_four_state
        )
        stage_six_state = self._stage_six_completion_state(
            project_path, checkpoints, targets, stage_one_state, stage_four_state, stage_five_state
        )

        project_overview_ready = stage_one_state["project_overview_ready"]
        notes_ready = stage_one_state["notes_ready"]
        references_ready = stage_one_state["references_ready"]
        outline_ready = stage_one_state["outline_ready"]
        research_plan_ready = stage_one_state["research_plan_ready"]

        data_log_quality_ok = stage_four_state["data_log_quality_ok"]
        analysis_quality_ok = stage_four_state["analysis_quality_ok"]
        report_ready = stage_four_state["report_ready"]
        independent_review_ready = stage_five_state["independent_review_ready"]
        lint_report_ready = stage_five_state["lint_report_ready"]
        review_reports_ready = stage_five_state["review_reports_ready"]
        presentation_ready = stage_six_state["presentation_ready"]
        delivery_ready = self._has_effective_delivery_log(project_path)
        presentation_required = stage_six_state["presentation_required"]

        interview_done = stage_one_state["interview_done"]
        outline_confirmed = stage_one_state["outline_confirmed"]
        review_started = stage_four_state["review_started"]
        review_passed = stage_five_state["review_passed"]
        presentation_done = stage_six_state["presentation_done"]
        delivery_archived = "delivery_archived_at" in checkpoints

        stage_zero_complete = stage_one_state["stage_zero_complete"]
        stage_one_complete = stage_one_state["stage_one_complete"]
        stage_two_complete = stage_four_state["stage_two_complete"]
        stage_three_complete = stage_four_state["stage_three_complete"]
        stage_four_complete = stage_four_state["stage_four_complete"]
        stage_five_complete = stage_five_state["stage_five_complete"]
        stage_six_complete = stage_six_state["stage_six_complete"]
        stage_seven_complete = stage_six_complete and delivery_ready and delivery_archived

        if not stage_zero_complete:
            stage_code = "S0"
            stage_status = "进行中"
        elif not stage_one_complete:
            stage_code = "S1"
            stage_status = "进行中"
        elif not stage_two_complete:
            stage_code = "S2"
            stage_status = "进行中"
        elif not stage_three_complete:
            stage_code = "S3"
            stage_status = "进行中"
        elif not stage_four_complete:
            stage_code = "S4"
            stage_status = "进行中"
        elif not stage_five_complete:
            stage_code = "S5"
            stage_status = "进行中"
        elif presentation_required and not stage_six_complete:
            stage_code = "S6"
            stage_status = "进行中"
        elif not stage_seven_complete:
            stage_code = "S7"
            stage_status = "进行中"
        else:
            stage_code = "done"
            stage_status = "已归档"

        # *_ready means effective file content; *_confirmed/started/passed/done/archived means a user checkpoint.
        flags = {
            "project_overview_ready": project_overview_ready,
            "s0_interview_done": interview_done,
            "notes_ready": notes_ready,
            "references_ready": references_ready,
            "outline_ready": outline_ready,
            "research_plan_ready": research_plan_ready,
            "data_log_ready": data_log_quality_ok,
            "analysis_ready": analysis_quality_ok,
            "report_ready": report_ready,
            "review_checklist_ready": False,
            "independent_review_ready": independent_review_ready,
            "lint_report_ready": lint_report_ready,
            "review_reports_ready": review_reports_ready,
            "review_notes_ready": self._has_effective_review_notes(project_path),
            "review_ready": review_reports_ready and review_passed,
            "presentation_ready": presentation_ready,
            "delivery_ready": delivery_ready and delivery_archived,
            "presentation_required": presentation_required,
            "outline_confirmed": outline_confirmed,
            "review_started": review_started,
            "review_passed": review_passed,
            "presentation_done": presentation_done,
            "delivery_archived": delivery_archived,
            "methodology_declared": self._methodology_declared_flag(project_path),
        }
        return {
            "stage_code": stage_code,
            "stage_status": stage_status,
            "completed_items": self._build_completed_items(stage_code, flags),
            "skipped_items": self._build_skipped_items(stage_code, flags),
            "checkpoints": checkpoints,
            "length_targets": targets,
            "flags": flags,
        }

    def _build_completed_items(self, stage_code: str, flags: dict) -> list[str]:
        completed: list[str] = []
        stage_index = self._stage_index(stage_code)
        for stage in self.STAGE_ORDER[:stage_index]:
            if stage == "S6" and not flags["presentation_required"]:
                continue
            if stage == "S5":
                if flags["independent_review_ready"]:
                    completed.append(self.STAGE_CHECKLIST_ITEMS["S5"][0])
                if flags["lint_report_ready"]:
                    completed.append(self.STAGE_CHECKLIST_ITEMS["S5"][1])
                completed.append(self.STAGE_CHECKLIST_ITEMS["S5"][2])
                continue
            completed.extend(self.STAGE_CHECKLIST_ITEMS[stage])

        if stage_code == "S0":
            if flags["project_overview_ready"]:
                completed.append(self.STAGE_CHECKLIST_ITEMS["S0"][2])
        elif stage_code == "S1":
            if flags["notes_ready"]:
                completed.append(self.STAGE_CHECKLIST_ITEMS["S1"][0])
            if flags["references_ready"]:
                completed.append(self.STAGE_CHECKLIST_ITEMS["S1"][1])
            # R5: 「分析框架确定」镜像方法论声明 parsed（display-only，不驱动阶段回归）
            if flags.get("methodology_declared") and flags["outline_ready"]:
                completed.append(self.STAGE_CHECKLIST_ITEMS["S1"][2])
            if flags["outline_ready"]:
                completed.append(self.STAGE_CHECKLIST_ITEMS["S1"][3])
            if flags["research_plan_ready"]:
                completed.append(self.STAGE_CHECKLIST_ITEMS["S1"][4])
        elif stage_code == "S3" and flags["analysis_ready"]:
            completed.extend(self.STAGE_CHECKLIST_ITEMS["S3"])
        elif stage_code == "S4" and flags["report_ready"]:
            completed.extend(self.STAGE_CHECKLIST_ITEMS["S4"])
        elif stage_code == "S5":
            if flags["independent_review_ready"]:
                completed.append(self.STAGE_CHECKLIST_ITEMS["S5"][0])
            if flags["lint_report_ready"]:
                completed.append(self.STAGE_CHECKLIST_ITEMS["S5"][1])
            if flags["independent_review_ready"] and flags["lint_report_ready"]:
                completed.append(self.STAGE_CHECKLIST_ITEMS["S5"][2])
        elif stage_code == "S6":
            completed.append(self.STAGE_CHECKLIST_ITEMS["S6"][0])
            if flags["presentation_ready"]:
                completed.extend(self.STAGE_CHECKLIST_ITEMS["S6"][1:])
        elif stage_code == "S7" and flags["delivery_ready"]:
            if flags["presentation_required"] and flags["presentation_ready"]:
                completed.extend(self.STAGE_CHECKLIST_ITEMS["S6"])
            completed.extend(self.STAGE_CHECKLIST_ITEMS["S7"])

        return list(dict.fromkeys(completed))

    def _build_skipped_items(self, stage_code: str, flags: dict) -> list[str]:
        if not flags["presentation_required"] and stage_code in {"S7", "done"}:
            return list(self.STAGE_CHECKLIST_ITEMS["S6"])
        return []

    def _stage_index(self, stage_code: str) -> int:
        if stage_code == "done":
            return len(self.STAGE_ORDER)
        return self.STAGE_ORDER.index(stage_code)

    def _tracked_stage_items(self) -> list[str]:
        tracked_items: list[str] = []
        for stage in self.STAGE_ORDER:
            tracked_items.extend(self.STAGE_CHECKLIST_ITEMS[stage])
        return tracked_items

    def _set_stage_gate_item_state(self, content: str, task: str, state: str) -> str:
        pattern = rf"- \[(?: |x|X|/)\] {re.escape(task)}"
        return re.sub(pattern, f"- [{state}] {task}", content)

    def _is_plan_markdown_path(self, normalized_path: str) -> bool:
        candidate = self._to_posix(normalized_path).lstrip("/").lower()
        return candidate.startswith("plan/") and candidate.endswith(".md")

    def _requires_pre_outline_evidence(self, normalized_path: str) -> bool:
        return self._canonicalize_plan_markdown_path(normalized_path) in {
            "plan/outline.md",
            "plan/research-plan.md",
        }

    def _is_backend_owned_stage_tracking_file(self, normalized_path: str) -> bool:
        return self._canonicalize_plan_markdown_path(normalized_path) in {
            "plan/stage-gates.md",
            "plan/progress.md",
            "plan/tasks.md",
        }

    def _canonicalize_plan_markdown_path(self, normalized_path: str) -> str:
        candidate = self._to_posix(normalized_path).lstrip("/")
        if not self._is_plan_markdown_path(candidate):
            return candidate
        return candidate.lower()

    def _is_effective_plan_file(self, project_path: Path, file_name: str) -> bool:
        text = self._read_plan_file(project_path, file_name)
        if not text:
            return False
        if self._is_template_content(text, file_name):
            return False
        return self._has_substantive_body(text)

    def _read_plan_file(self, project_path: Path, file_name: str) -> str:
        file_path = project_path / "plan" / file_name
        if not file_path.exists():
            return ""
        try:
            return file_path.read_text(encoding="utf-8").strip()
        except (OSError, UnicodeDecodeError):
            return ""

    def _read_project_file(self, project_path: Path, relative_path: str) -> str:
        file_path = project_path / relative_path
        if not file_path.exists():
            return ""
        return file_path.read_text(encoding="utf-8").strip()

    def _is_template_content(self, text: str, file_name: str) -> bool:
        template_path = self.skill_dir / "plan-template" / file_name
        if not template_path.exists():
            return False
        template_text = template_path.read_text(encoding="utf-8").strip()
        return self._normalize_text(text) == self._normalize_text(template_text)

    def _has_substantive_body(self, text: str) -> bool:
        return self._count_substantive_body_lines(text) > 0

    def _count_substantive_body_lines(self, text: str) -> int:
        count = 0
        for raw_line in text.splitlines():
            stripped = raw_line.strip()
            if not stripped or stripped == "---":
                continue
            if stripped.startswith("#"):
                continue
            if stripped.startswith("<!--") and stripped.endswith("-->"):
                continue
            if stripped.startswith("|"):
                continue
            candidate = stripped[2:].strip() if stripped.startswith("- ") else stripped
            candidate = re.sub(r"\*\*([^*]+)\*\*\s*[:：]\s*", "", candidate).strip()
            if self._is_substantive_field_value(candidate):
                count += 1
        return count

    def _is_substantive_field_value(self, value: str) -> bool:
        candidate = re.sub(r"\*\*", "", (value or "").strip())
        if not candidate or candidate in {"-", "|"}:
            return False
        if re.search(r"\[[ xX/]\]", candidate):
            return False
        if candidate.startswith("[") and candidate.endswith("]"):
            return False
        if re.fullmatch(r"\d+\.", candidate):
            return False
        if re.fullmatch(r"[|\-:\s]+", candidate):
            return False
        if candidate.endswith(":") or candidate.endswith("："):
            return False
        return True

    def _has_effective_notes(self, project_path: Path) -> bool:
        notes_text = self._read_plan_file(project_path, "notes.md")
        if not notes_text or self._is_template_content(notes_text, "notes.md"):
            return False
        heading_count = len(re.findall(r"^(?:##+|###)\s+", notes_text, flags=re.MULTILINE))
        labeled_line_count = self._count_substantive_labeled_lines(notes_text)
        substantive_line_count = self._count_substantive_body_lines(notes_text)
        return substantive_line_count >= 2 and (heading_count >= 2 or labeled_line_count >= 2)

    def _has_effective_references(self, project_path: Path) -> bool:
        references_text = self._read_plan_file(project_path, "references.md")
        if not references_text or self._is_template_content(references_text, "references.md"):
            return False
        return self._count_reference_evidence(references_text) >= 2

    def _count_reference_evidence(self, references_text: str) -> int:
        count = 0
        for raw_line in references_text.splitlines():
            line = raw_line.strip()
            if not line or line.startswith("#") or line == "---":
                continue
            if line.startswith("- "):
                entry = line[2:].strip()
                if self._looks_like_reference_evidence(entry):
                    count += 1
                    continue
            numbered_match = re.match(r"^\d+\.\s+(.+)$", line)
            if numbered_match:
                entry = numbered_match.group(1).strip()
                if self._looks_like_reference_evidence(entry):
                    count += 1
                    continue
            if line.startswith("**") and ":" in line:
                _, value = line.split(":", 1)
                value = value.strip()
                if self._looks_like_reference_evidence(value):
                    count += 1
                    continue
            if "http://" in line or "https://" in line:
                count += 1
        return count

    def _normalize_reference_evidence_value(self, value: str) -> str:
        candidate = re.sub(r"\[([^\]]+)\]\((https?://[^)]+)\)", r"\1 \2", (value or "").strip())
        return re.sub(r"^\[[^\]]+\]\s*", "", candidate).strip()

    def _looks_like_reference_evidence(self, value: str) -> bool:
        candidate = self._normalize_reference_evidence_value(value)
        if not self._is_substantive_field_value(candidate):
            return False

        lowered = candidate.lower()
        placeholder_tokens = (
            "tbd",
            "todo",
            "placeholder",
            "source name",
            "待补",
            "待确认",
            "来源名称",
            "示例来源",
        )
        if any(token in lowered for token in placeholder_tokens):
            return False

        if "[" in candidate or "]" in candidate:
            return False

        return True

    def _has_effective_outline(self, project_path: Path) -> bool:
        outline_text = self._read_plan_file(project_path, "outline.md")
        if not outline_text or self._is_template_content(outline_text, "outline.md"):
            return False
        section_count = len(re.findall(r"^(?:##+\s+|[0-9]+\.\s+)", outline_text, flags=re.MULTILINE))
        return section_count >= 2 and self._has_substantive_body(outline_text)

    def _has_effective_research_plan(self, project_path: Path) -> bool:
        research_plan_text = self._read_plan_file(project_path, "research-plan.md")
        if not research_plan_text or self._is_template_content(research_plan_text, "research-plan.md"):
            return False
        section_count = len(
            re.findall(r"^(?:##+\s+|[0-9]+\.\s+)", research_plan_text, flags=re.MULTILINE)
        )
        required_patterns = [
            r"research methods?",
            r"data sources?",
            r"execution steps?",
            r"研究方法",
            r"数据来源",
            r"执行步骤",
        ]
        structural_patterns = [
            r"research objectives?",
            r"research questions?",
            r"phase plan",
            r"key assumptions?",
            r"研究目标",
            r"研究问题",
            r"核心研究问题",
            r"阶段安排",
            r"关键假设",
        ]
        has_named_sections = any(
            re.search(pattern, research_plan_text, flags=re.IGNORECASE)
            for pattern in required_patterns
        )
        has_plan_structure = any(
            re.search(pattern, research_plan_text, flags=re.IGNORECASE)
            for pattern in structural_patterns
        )
        return (has_named_sections or (section_count >= 2 and has_plan_structure)) and self._has_substantive_body(
            research_plan_text
        )

    def _has_effective_data_log(self, project_path: Path) -> bool:
        data_log_text = self._read_plan_file(project_path, "data-log.md")
        if not data_log_text or self._is_template_content(data_log_text, "data-log.md"):
            return False
        seen_table_separator = False
        for raw_line in data_log_text.splitlines():
            line = raw_line.strip()
            if not line or line.startswith("#"):
                continue
            if line.startswith("|"):
                cells = [cell.strip() for cell in line.strip("|").split("|")]
                if len(cells) < 4:
                    continue
                if re.fullmatch(r"[-:\s]+", "".join(cells)):
                    seen_table_separator = True
                    continue
                if not seen_table_separator:
                    continue
                if self._is_substantive_field_value(cells[3]) and any(
                    self._is_substantive_field_value(cell) for cell in cells[:3]
                ):
                    return True
                continue
            if line.startswith("- ") and self._is_substantive_field_value(line[2:].strip()):
                return True
        return False

    def _has_effective_analysis_notes(self, project_path: Path) -> bool:
        analysis_text = self._read_plan_file(project_path, "analysis-notes.md")
        if not analysis_text or self._is_template_content(analysis_text, "analysis-notes.md"):
            return False
        insight_heading_count = len(re.findall(r"^(?:##+)\s+", analysis_text, flags=re.MULTILINE))
        labeled_line_count = self._count_substantive_labeled_lines(analysis_text)
        return insight_heading_count >= 1 and labeled_line_count >= 3 and self._has_substantive_body(analysis_text)

    def _has_effective_review_checklist(self, project_path: Path) -> bool:
        review_text = self._read_plan_file(project_path, "review-checklist.md")
        if not review_text or self._is_template_content(review_text, "review-checklist.md"):
            return False
        total_items = re.findall(r"^\s*-\s+\[[ xX/]\]\s+.+\S$", review_text, flags=re.MULTILINE)
        checked_items = re.findall(r"^\s*-\s+\[[xX]\]\s+.+\S$", review_text, flags=re.MULTILINE)
        return len(total_items) >= 3 and len(checked_items) == len(total_items)

    def _has_effective_independent_review(self, project_path: Path) -> bool:
        review_text = self._read_plan_file(project_path, "independent-review.md")
        if not review_text or self._is_template_content(review_text, "independent-review.md"):
            return False
        if not all(anchor in review_text for anchor in self.INDEPENDENT_REVIEW_ANCHORS):
            return False
        if self.INDEPENDENT_REVIEW_COMPLETION_MARKER not in review_text:
            return False
        return self._has_substantive_body(review_text)

    def _has_effective_lint_report(self, project_path: Path) -> bool:
        lint_text = self._read_plan_file(project_path, "lint-report.md")
        if not lint_text or self._is_template_content(lint_text, "lint-report.md"):
            return False
        if not all(anchor in lint_text for anchor in self.LINT_REPORT_ANCHORS):
            return False
        if self.LINT_REPORT_COMPLETION_MARKER not in lint_text:
            return False
        return self._has_substantive_body(lint_text)

    def _has_effective_review_reports(self, project_path: Path) -> bool:
        return (
            self._has_effective_independent_review(project_path)
            and self._has_effective_lint_report(project_path)
        )

    def _is_report_review_stale(self, project_path: Path) -> bool:
        """R3 D6 advisory: both review reports are EFFECTIVE (substantive, not the scaffolded
        template — BLOCKER 1) AND the draft is newer than the OLDER report. NOT gated on
        review_passed_at — covers the window where reports exist, the draft was edited, but the
        user hasn't clicked 审查通过 yet (review_passed_at unset; record_stage_checkpoint only
        checks report structure, not whether they cover the current draft)."""
        draft_path = project_path / self.REPORT_DRAFT_PATH
        if not draft_path.exists():
            return False
        # Templates only (new-project scaffold of independent-review.md / lint-report.md) don't
        # count — both must be effective reports, reusing the production gate.
        if not self._has_effective_review_reports(project_path):
            return False
        ir_path = project_path / "plan" / "independent-review.md"
        lint_path = project_path / "plan" / "lint-report.md"
        draft_mtime = draft_path.stat().st_mtime_ns
        oldest_report_mtime = min(ir_path.stat().st_mtime_ns, lint_path.stat().st_mtime_ns)
        return draft_mtime > oldest_report_mtime

    def _has_effective_review_notes(self, project_path: Path) -> bool:
        review_text = self._read_plan_file(project_path, "review.md")
        if not review_text or self._is_template_content(review_text, "review.md"):
            return False
        substantive_review_item = re.search(
            r"^\s*(?:-\s+(?!\[[ xX/]\])[^\s].*|\d+\.[ \t]+[^\s].*)$",
            review_text,
            flags=re.MULTILINE,
        )
        substantive_labeled_line = self._count_substantive_labeled_lines(review_text) >= 2
        return bool(substantive_review_item or substantive_labeled_line) and self._has_substantive_body(review_text)

    def _count_substantive_labeled_lines(self, text: str) -> int:
        count = 0
        plain_text = text.replace("**", "")
        for raw_line in plain_text.splitlines():
            line = raw_line.strip()
            if not line or line.startswith("#"):
                continue
            match = re.search(r"^[^:：]+[:：]\s*(.+)$", line)
            if match and self._is_substantive_field_value(match.group(1)):
                count += 1
        return count

    def _has_effective_presentation_plan(self, project_path: Path) -> bool:
        presentation_text = self._read_plan_file(project_path, "presentation-plan.md")
        if not presentation_text or self._is_template_content(presentation_text, "presentation-plan.md"):
            return False
        required_patterns = [r"\bppt\b", r"\bq&a\b", r"\bnarrative\b", r"演示", r"讲稿"]
        return (
            any(re.search(pattern, presentation_text, flags=re.IGNORECASE) for pattern in required_patterns)
            and self._has_substantive_body(presentation_text)
        )

    def _has_effective_delivery_log(self, project_path: Path) -> bool:
        delivery_text = self._read_plan_file(project_path, "delivery-log.md")
        if not delivery_text or self._is_template_content(delivery_text, "delivery-log.md"):
            return False
        required_patterns = [r"\bdelivery date\b", r"\bshared\b", r"\bsent\b", r"\bfeedback\b", r"交付", r"反馈"]
        return (
            any(re.search(pattern, delivery_text, flags=re.IGNORECASE) for pattern in required_patterns)
            and self._has_substantive_body(delivery_text)
        )

    def _count_words(self, content: str) -> int:
        text = content
        for pattern, repl in self._MARKDOWN_STRIP_PATTERNS:
            text = pattern.sub(repl, text)
        stripped = re.sub(r"[\s\u3000]+", "", text)
        return len(stripped)

    def _is_template_stub(self, text: str) -> bool:
        return not self._has_substantive_body(text)

    def _has_effective_report_draft(self, project_path: Path, min_words: int = 0) -> bool:
        draft_text = self._read_project_file(project_path, self.REPORT_DRAFT_PATH)
        if not draft_text or self._is_template_stub(draft_text):
            return False
        if min_words and self._count_words(draft_text) < min_words:
            return False
        return True

    def _delivery_mode_requires_presentation(self, project_path: Path) -> bool:
        return self._extract_delivery_mode(project_path) == "报告+演示"

    def load_type_skeleton(self, project_type: str) -> str:
        """取类型模块的「## 二、标准结构」段作为报告骨架。caller 保证 project_type ∈
        TYPE_SKELETON_MAP（未知 type 在 build_methodology_block 已 graceful 返空）。
        已知 type 但模块缺锚点 / 段为空 → fail-closed 抛 ValueError（代码/资产回归立刻暴露）。
        逐行扫描并跳过 ``` 代码块，避免被骨架代码块内的 `## 执行摘要` 等行提前截断。"""
        filename = self.TYPE_SKELETON_MAP[project_type]
        module_path = self.skill_dir / "modules" / filename
        try:
            text = module_path.read_text(encoding="utf-8")
        except FileNotFoundError as exc:
            raise ValueError(f"模块 {filename} 不存在（已知 type 但骨架资产缺失）") from exc
        lines = text.splitlines()
        anchor_idx = None
        for idx, line in enumerate(lines):
            if re.match(r"^##\s*二、标准结构\s*$", line):
                anchor_idx = idx
                break
        if anchor_idx is None:
            raise ValueError(f"模块 {filename} 缺少「## 二、标准结构」锚点")
        body_lines = []
        in_fence = False
        for line in lines[anchor_idx + 1:]:
            if line.lstrip().startswith("```"):
                in_fence = not in_fence
                body_lines.append(line)
                continue
            if not in_fence and re.match(r"^##\s", line):
                break
            body_lines.append(line)
        if in_fence:
            raise ValueError(f"模块 {filename}「## 二、标准结构」段有未闭合代码块（``` 不成对）")
        body = "\n".join(body_lines).strip()
        if not body:
            raise ValueError(f"模块 {filename}「## 二、标准结构」段为空")
        return body

    @staticmethod
    def _normalize_for_danger(text: str) -> str:
        """归一化用于危险词比对：NFKC + casefold + 删 Unicode 格式字符（Cf：零宽空格 U+200B/
        BOM U+FEFF/零宽连接符等）+ 去所有空白与常见分隔符。
        不变式（防拆词绕过，红队 v2/v4）：去除集合必须 ⊇ parse 的 split 分隔符（、,，）∪ off-menu
        白名单允许的非字母数字字符（- / 空格 全角空格）——这样任何被允许字符拆开的 API 名（如
        「write、file」「advance stage」「s0-interview-done-at」）归一化后都还原成连续串、命中
        denylist。改 off-menu 白名单或 split 分隔符时必须同步本集合。"""
        folded = unicodedata.normalize("NFKC", text or "").casefold()
        folded = "".join(ch for ch in folded if unicodedata.category(ch) != "Cf")
        return re.sub(r"[\s\-_/.·、,，]", "", folded)

    def _canonical_framework_name(self, token: str) -> Optional[str]:
        """token 去空格 casefold 后命中已知框架名 → 返回去空白原文；否则 None。
        归一化让「BCG 矩阵」「ISO 8000」「DAMA-DMBOK / ISO 8000」等带空格写法也命中。"""
        normalized = token.casefold().replace(" ", "").replace("　", "")
        if normalized in self.KNOWN_FRAMEWORK_NAMES:
            return token.strip()
        return None

    def parse_and_sanitize_methodology(self, outline_text: str) -> tuple[str, list[str]]:
        """解析 outline 顶部「方法论框架：…」声明行，净化为可信框架名列表。
        返回 (state, frameworks)：
          - parsed   : 至少一个合法框架（已知名精确匹配，或菜单外严格短标签）
          - missing  : 顶部无声明行（legacy / 漏写）
          - malformed: 顶部有声明行但含工具名/checkpoint/注入词/阶段操控词，或全是非法标签
        净化是 trust boundary（outline 用户可编辑，spec §4.2/§11）：危险词在完整 raw_value 上
        先行检测（原样子串 + 归一化子串双查，挡分隔符绕过）、命中即整条 malformed（不剥括号、
        不截断绕过）；净化结果以「数据」注入，绝不当指令。"""
        text = outline_text or ""
        # 仅解析「顶部」声明（spec §7.1）：首个二级及以上标题（## ~ ######，允许 ≤3 前导空格）
        # 之前、且不超过前 30 行。H1（# 报告大纲）不截断（声明在 H1 之后）；缩进 H2 也截断（防把
        # 真声明挤出顶部，红队 v2）。避免正文/示例/代码块里的「方法论框架：」被误解析。
        head_lines: list[str] = []
        for line in text.splitlines()[:30]:
            if re.match(r"^[^\S\n]{0,3}#{2,6}[^\S\n]", line):
                break
            head_lines.append(line)
        head = "\n".join(head_lines)
        match = self._METHODOLOGY_DECLARATION_RE.search(head)
        if not match:
            return ("missing", [])
        raw_value = match.group(1).strip()
        # 危险词先行（完整 raw_value，未剥括号、未截断）——spec §11「含危险词→malformed，不剥」。
        # 原样 casefold 子串挡注入符号/中文操控词；归一化（去空白+分隔符）子串挡工具名/checkpoint
        # 的「advance stage」「advance-stage」等分隔符变体（红队 v2）。
        lowered_raw = raw_value.casefold()
        normalized_raw = self._normalize_for_danger(raw_value)
        if any(bad in lowered_raw for bad in self._METHODOLOGY_DANGER_SUBSTRINGS) or any(
            bad in normalized_raw for bad in self._METHODOLOGY_DANGER_NORMALIZED
        ):
            return ("malformed", [])
        # 仅顿号/中英逗号分隔；不用 "/"（TAM-SAM-SOM、BCG/GE、金字塔原理/MECE 内部含 "/"）。
        tokens = [t.strip() for t in re.split(r"[、,，]+", raw_value) if t.strip()]
        if not tokens:
            return ("malformed", [])
        cleaned: list[str] = []
        for token in tokens[:8]:  # 条数上限（危险词已在 raw_value 层全量检测，截断不漏检）
            bare = re.sub(r"[（(].*?[)）]", "", token).strip()  # 剥括号（仅展示清洗，危险词已先拦）
            if not bare:
                continue
            canonical = self._canonical_framework_name(bare)
            if canonical:  # 已知框架名精确放行
                cleaned.append(canonical)
                continue
            # 菜单外：严格短标签（中英文/数字/连字符/斜杠/空格，≤24 字；工具名/checkpoint/注入词
            # 已被 raw_value 层双查拦截）。允许空格让「麦肯锡 7S」等带空格框架不被误杀。
            if re.fullmatch(r"[A-Za-z0-9一-鿿\-/ 　]{1,24}", bare):
                cleaned.append(bare)
                continue
            return ("malformed", [])
        if not cleaned:
            return ("malformed", [])
        # 归一化去重（去空白+分隔符）：「TAM-SAM-SOM」与「TAM SAM SOM」合并，保留首个 display。
        deduped: list[str] = []
        seen: set[str] = set()
        for name in cleaned:
            key = self._normalize_for_danger(name)
            if key not in seen:
                seen.add(key)
                deduped.append(name)
        return ("parsed", deduped)

    def _get_project_type_for_path(self, project_path: Path) -> Optional[str]:
        """按 project_dir 反查 project_type（registry 字段）。build_methodology_block 有
        project_id 直接取；确认门/快照只有 project_path，用本 helper 反查（门禁/注入同源口径）。"""
        try:
            target = Path(project_path).resolve()
        except OSError:
            target = Path(project_path)
        for project in self._load_registry()["projects"]:
            project_dir = project.get("project_dir")
            if not project_dir:
                continue
            try:
                if Path(project_dir).resolve() == target:
                    return project.get("project_type")
            except OSError:
                continue
        return None

    def read_confirmed_methodology_snapshot(self, project_path) -> tuple[str, list[str]]:
        """读「确认大纲那刻」冻结的方法论快照（非活 outline，跨轮/跨压缩稳定）。
        返回 (parsed/missing, frameworks)。快照仅在 _snapshot_methodology_on_confirm 解析为
        parsed 时写入（malformed/missing 不写），故读取只有 parsed/missing 两态。"""
        raw = self._read_raw_stage_checkpoints(project_path)
        snapshot = raw.get(self.METHODOLOGY_SNAPSHOT_KEY)
        if not isinstance(snapshot, str) or not snapshot.strip():
            return ("missing", [])
        frameworks = [token.strip() for token in snapshot.split("、") if token.strip()]
        if not frameworks:
            return ("missing", [])
        return ("parsed", frameworks)

    def _snapshot_methodology_on_confirm(self, project_path: Path) -> None:
        """确认大纲那刻：解析+净化 outline 声明，冻结进 __methodology_snapshot 保留键。
        未知 type / 无有效声明 → 不写（S2–S4 注入靠 read_confirmed_methodology_snapshot 的
        missing 兜底）。后端写、非模型写、非新 checkpoint key。"""
        project_type = self._get_project_type_for_path(project_path)
        if project_type not in self.TYPE_SKELETON_MAP:
            return
        outline_text = self._read_plan_file(project_path, "outline.md") or ""
        state, selected = self.parse_and_sanitize_methodology(outline_text)
        if state != "parsed" or not selected:
            return
        raw = self._read_raw_stage_checkpoints(project_path)
        raw[self.METHODOLOGY_SNAPSHOT_KEY] = "、".join(selected)
        self._write_raw_stage_checkpoints(project_path, raw)

    def get_project_type(self, project_ref: str) -> Optional[str]:
        record = self.get_project_record(project_ref)
        return record.get("project_type") if record else None

    def _declare_and_invite_instruction(self, project_type: str) -> str:
        """S1 注入：让模型在 outline 顶部写方法论声明行 + 聊天里软邀请（按类型分腔调，§7.3）。"""
        tone = self.METHODOLOGY_TONE.get(project_type, "analytical")
        # 注意：腔调举例里框架之间一律用「顿号」分隔，与声明行格式（顿号分隔）一致——
        # 否则模型照提示用 + / 空格连接，会被 B3 parser 判 malformed、卡住确认门（codex R1 BLOCKER 4）。
        if tone == "bid":
            tone_line = (
                "本技术标的方法＝依招标文件/技规评分点组织结构，并逐条响应；"
                "在声明里写清所用方法（如评分点对标、点对点应答、WBS、重难点对策）。"
                "结构以招标文件为准，不要硬贴通用分析框架。"
            )
        elif tone == "structural":
            tone_line = (
                "本报告的「方法论」是结构纪律：管理制度用「章-条-款-项」规范结构；"
                "实施方案用 SMART、RACI、里程碑。按本报告类型选，不要硬贴 SWOT 之类分析框架。"
            )
        elif tone == "specialized":
            tone_line = (
                "按本专项研究的子题目选方法：数据治理题用 DAMA-DMBOK、ISO 8000、成熟度模型；"
                "非数据题用根因分析、对标分析，不要硬套招牌框架。"
            )
        else:  # analytical
            tone_line = (
                "从下方框架菜单挑本报告真正需要的招牌框架（如 SWOT、波特五力、BCG 矩阵），"
                "也可以用你自己知道的其他框架。"
            )
        return (
            "## 方法论声明（S1）\n"
            f"{tone_line}\n"
            "在 `plan/outline.md` 第一行（在 `## 确认状态` 等二级标题之前）写一行可见声明（格式固定，供系统识别）：\n"
            "`方法论框架：〔框架1〕、〔框架2〕`（顿号分隔，可加粗 `**方法论框架**：…`）。\n"
            "写完声明后，在聊天里顺口告诉用户本报告将采用〔所选框架〕；若用户想换方法论，"
            "告诉你即可，否则按这个继续，可随时在工作区点「确认大纲」。"
        )

    def _adhere_instruction(self, state: str, selected: list[str]) -> str:
        """S2–S4 注入：沿用确认时快照的已选框架，不再邀请重选。malformed 不入快照，故只两态。"""
        if state == "parsed" and selected:
            joined = "、".join(selected)
            return (
                "## 方法论（已选）\n"
                f"本报告已选方法论框架：{joined}。正文须沿用，不要重新征求或反复改大纲方法论。"
                "如用户要大改方法论，提示需回 S1 调整大纲并重新确认。"
            )
        return (
            "## 方法论\n"
            "本报告未记录已确认的方法论框架。按报告类型与框架菜单选合适框架展开分析，"
            "保持结论先行、结构清晰；不要凭空声称某框架是「已确认」的。"
        )

    def _framework_menu_for_type(self, project_type: str) -> str:
        """技术标按评分点驱动、逐条响应，不靠「挑分析框架」；通用框架菜单对它既误导又
        挤爆 token 预算（spec §3.2，2026-06-20 用户拍板）→ bid 不注入菜单。其余类型沿用。"""
        if project_type == "technical-bid":
            return ""
        return self.FRAMEWORK_MENU

    def _render_methodology_block(self, skeleton: str, menu: str, instr: str) -> str:
        return (
            "# 方法论与报告结构（系统按报告类型注入）\n\n"
            "## 报告结构骨架（按类型）\n"
            f"{skeleton}\n\n"
            f"{menu}\n"
            f"{instr}"
        )

    def build_methodology_block(self, project_id: str) -> str:
        """按 project_type 注入「类型骨架 + 框架菜单 + 阶段化指令」到 system prompt（S1–S4）。
        装配期只读，不写任何文件。未知 type / 非写作期 → graceful 空块（绝不抛进 chat 链路，
        codex R2 BLOCKER 5）；已知 type 但模块缺锚点 → load_type_skeleton fail-closed 抛（§4.1）。"""
        project_path = self.get_project_path(project_id)
        if project_path is None:
            return ""
        stage = self._infer_stage_state(project_path)["stage_code"]
        if stage not in ("S1", "S2", "S3", "S4"):
            return ""
        project_type = self.get_project_type(project_id)
        if project_type not in self.TYPE_SKELETON_MAP:
            logger.info("unknown project_type %r, skip methodology block", project_type)
            return ""
        skeleton = self.load_type_skeleton(project_type)
        if stage == "S1":
            instr = self._declare_and_invite_instruction(project_type)
        else:
            state, selected = self.read_confirmed_methodology_snapshot(project_path)
            instr = self._adhere_instruction(state, selected)
        return self._render_methodology_block(
            skeleton, self._framework_menu_for_type(project_type), instr
        )

    def _methodology_declared_flag(self, project_path: Path) -> bool:
        """前端确认按钮用：known type + 未确认时，要求 outline 有 parsed 声明才 True；
        unknown type / 已确认 → True（不门禁 / 不再卡）。仅 known+未确认时有约束意义。"""
        project_type = self._get_project_type_for_path(project_path)
        if project_type not in self.TYPE_SKELETON_MAP:
            return True
        checkpoints = self._load_stage_checkpoints(project_path)
        if "outline_confirmed_at" in checkpoints:
            return True
        outline_text = self._read_plan_file(project_path, "outline.md") or ""
        state, _ = self.parse_and_sanitize_methodology(outline_text)
        return state == "parsed"

    def get_skill_prompt(self) -> str:
        """鑾峰彇Skill瀹氫箟"""
        skill_file = self.skill_dir / "SKILL.md"
        sections = [skill_file.read_text(encoding="utf-8")]

        lifecycle_file = self.skill_dir / "modules" / "consulting-lifecycle.md"
        if lifecycle_file.exists():
            sections.append("## Consulting Lifecycle Guidance\n" + lifecycle_file.read_text(encoding="utf-8"))

        return "\n\n".join(sections)

    def _normalize_create_payload(self, project_info_or_name, **kwargs) -> dict:
        if hasattr(project_info_or_name, "model_dump"):
            payload = project_info_or_name.model_dump()
        elif isinstance(project_info_or_name, dict):
            payload = dict(project_info_or_name)
        elif project_info_or_name is None:
            payload = dict(kwargs)
        else:
            payload = {"name": project_info_or_name, **kwargs}

        payload.setdefault("notes", "")
        # 目标读者已非必填：缺省 / 显式 None / 纯空白一律归一为 ""，
        # 否则 None 会流到 _populate_v2_plan_files 的 str.replace 触发 TypeError。
        payload["target_audience"] = str(payload.get("target_audience") or "").strip()
        payload.setdefault("workspace_dir", kwargs.get("workspace_dir"))
        payload.setdefault("initial_material_paths", kwargs.get("initial_material_paths") or [])
        required_fields = [
            "name",
            "project_type",
            "theme",
            "deadline",
            "expected_length",
        ]
        missing = [field for field in required_fields if not payload.get(field)]
        if missing:
            raise ValueError(f"缺少项目字段: {', '.join(missing)}")
        return payload

    def _load_registry(self) -> dict:
        if not self.registry_path.exists():
            return {"projects": []}
        return json.loads(self.registry_path.read_text(encoding="utf-8"))

    def _save_registry(self, registry: dict):
        self.registry_path.write_text(
            json.dumps(registry, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    def _touch_project(self, project_id: str):
        registry = self._load_registry()
        for project in registry["projects"]:
            if project["id"] == project_id:
                project["updated_at"] = datetime.now().isoformat(timespec="seconds")
                break
        self._save_registry(registry)

    def _materials_path(self, project_record: dict) -> Path:
        return Path(project_record["project_dir"]) / "materials.json"

    def _load_materials(self, project_record: dict) -> list[dict]:
        materials_path = self._materials_path(project_record)
        if not materials_path.exists():
            return []
        return json.loads(materials_path.read_text(encoding="utf-8"))

    def _save_materials(self, project_record: dict, materials: list[dict]):
        self._materials_path(project_record).write_text(
            json.dumps(materials, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    def _workspace_relative_path(self, source_path: Path, workspace_root: Path) -> Optional[str]:
        if self._is_within(source_path, workspace_root):
            return self._to_posix(source_path.relative_to(workspace_root))
        return None

    def _find_existing_workspace_material(self, materials: list[dict], stored_rel_path: str) -> Optional[dict]:
        return next(
            (
                item for item in materials
                if item["source_type"] == "workspace" and item["stored_rel_path"] == stored_rel_path
            ),
            None,
        )

    def _find_existing_imported_material(self, materials: list[dict], source_path: Path) -> Optional[dict]:
        source_str = str(source_path)
        return next(
            (
                item for item in materials
                if item["source_type"] == "imported" and item.get("original_path") == source_str
            ),
            None,
        )

    def _build_imported_destination(self, project_path: Path, file_name: str) -> Path:
        imported_root = project_path / "materials" / "imported"
        candidate_name = file_name
        stem = Path(file_name).stem
        suffix = Path(file_name).suffix
        counter = 1
        while (imported_root / candidate_name).exists():
            candidate_name = f"{stem}-{counter}{suffix}"
            counter += 1
        return Path("materials") / "imported" / candidate_name

    def _detect_media_kind(self, source_path: Path) -> str:
        return "image_like" if source_path.suffix.lower() in self.IMAGE_SUFFIXES else "text_like"

    def _new_id(self, prefix: str) -> str:
        return f"{prefix}-{uuid.uuid4().hex[:12]}"

    def _to_posix(self, path_value) -> str:
        if isinstance(path_value, Path):
            return path_value.as_posix()
        return str(path_value).replace("\\", "/")

    def _resolve_project_path(self, project_path: Path, file_path: str) -> Path:
        full_path = (project_path / file_path).resolve()
        if not self._is_within(full_path, project_path.resolve()):
            raise ValueError("非法的文件路径")
        return full_path

    def _is_within(self, path: Path, base: Path) -> bool:
        try:
            path.resolve().relative_to(base.resolve())
            return True
        except ValueError:
            return False

    def _read_optional(self, project_ref: str, file_path: str) -> str:
        try:
            return self.read_file(project_ref, file_path)
        except ValueError:
            return ""

    def _extract_stage_code(self, progress_text: str) -> str:
        match = re.search(r"\*\*阶段\*\*:\s*([A-Z]\d)", progress_text)
        return match.group(1) if match else ""

    def _extract_stage_status(self, progress_text: str) -> str:
        match = re.search(r"\*\*状态\*\*:\s*([^\n]+)", progress_text)
        return match.group(1).strip() if match else ""

    def _extract_checked_items(self, markdown_text: str) -> list[str]:
        return [
            match.group(1).strip()
            for match in re.finditer(r"- \[x\]\s+(.+)", markdown_text, flags=re.IGNORECASE)
        ]

    def _extract_open_items(self, markdown_text: str) -> list[str]:
        return [
            match.group(1).strip()
            for match in re.finditer(r"- \[ \]\s+(.+)", markdown_text)
        ]

    def _normalize_text(self, text: str) -> str:
        return re.sub(r"\s+", " ", text).strip()

    def _read_docx(self, material_path: Path) -> str:
        from docx import Document

        document = Document(material_path)
        lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))
        return "\n".join(lines).strip()

    def _read_xlsx(self, material_path: Path) -> str:
        import openpyxl

        workbook = openpyxl.load_workbook(material_path, data_only=True, read_only=True)
        sections = []
        for sheet in workbook.worksheets:
            rows = []
            for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                values = ["" if value is None else str(value) for value in row]
                if any(values):
                    rows.append(" | ".join(values))
                if row_index >= 50:
                    break
            sections.append(f"## {sheet.title}\n" + "\n".join(rows))
        return "\n\n".join(section for section in sections if section.strip()).strip()

    def _read_pdf(self, material_path: Path) -> str:
        from pypdf import PdfReader

        reader = PdfReader(str(material_path))
        sections = []
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                sections.append(f"## 第{index}页\n{text}")
        if sections:
            return "\n\n".join(sections)
        raise ValueError("PDF 未提取到文本，当前版本暂不支持扫描版 PDF。")
